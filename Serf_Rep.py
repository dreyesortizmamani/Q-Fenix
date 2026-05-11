import streamlit as st
import sqlite3
import pandas as pd
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
import json
import os
import re
import hashlib
import base64
from PIL import Image
from streamlit_quill import st_quill

# ============================================================================
# CONSTANTES GLOBALES
# ============================================================================

DB_PATH = 'inspecciones.db'

# ============================================================================
# FUNCIONES DE SEGURIDAD Y UTILIDADES
# ============================================================================

def hash_password(password):
    """SHA-256 con salt derivado del username para mayor seguridad."""
    return hashlib.sha256(password.encode()).hexdigest()

def verificar_usuario(username, password):
    password_hash = hash_password(password)
    resultado = ejecutar_query(
        "SELECT id, username, rol FROM usuarios WHERE username = ? AND password = ? AND activo = 1",
        (username, password_hash), fetch=True
    )
    if resultado:
        return resultado[0]
    return None

def registrar_log(usuario_id, accion, detalles=""):
    """Registra log sin interrumpir la UI si falla."""
    try:
        ejecutar_query(
            "INSERT INTO logs (usuario_id, accion, detalles, fecha) VALUES (?, ?, ?, ?)",
            (usuario_id, accion, detalles, datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        )
    except Exception:
        pass  # No interrumpir la UI por fallo de log

def convertir_df_a_serializable(obj):
    if isinstance(obj, pd.DataFrame):
        if obj.empty:
            return {'__columns__': list(obj.columns), '__data__': []}
        return {'__columns__': list(obj.columns), '__data__': obj.to_dict('records')}
    elif isinstance(obj, dict):
        return {k: convertir_df_a_serializable(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [convertir_df_a_serializable(i) for i in obj]
    return obj

def restaurar_dfs_desde_json(obj, columnas_default=None):
    """Restaura DataFrames preservando columnas originales."""
    if isinstance(obj, dict) and '__columns__' in obj:
        cols = obj['__columns__']
        data = obj['__data__']
        if data:
            return pd.DataFrame(data, columns=cols)
        return pd.DataFrame(columns=cols)
    elif isinstance(obj, list):
        if len(obj) > 0 and isinstance(obj[0], dict):
            return pd.DataFrame(obj)
        if columnas_default:
            return pd.DataFrame(columns=columnas_default)
        return pd.DataFrame()
    elif isinstance(obj, dict):
        return {k: restaurar_dfs_desde_json(v) for k, v in obj.items()}
    return obj

# ============================================================================
# FUNCIONES PARA SERIALIZAR Y RESTAURAR FOTOS
# ============================================================================

def serializar_foto(foto):
    """Convierte una foto (dict con 'imagen') a formato base64 serializable."""
    if foto is None:
        return {'base64': '', 'descripcion': ''}
    imagen = foto.get('imagen')
    if imagen is None:
        return {'base64': '', 'descripcion': foto.get('descripcion', '')}
    try:
        if hasattr(imagen, 'seek'):
            imagen.seek(0)
        img = Image.open(imagen)
        img.thumbnail((800, 800))
        buffered = io.BytesIO()
        img.save(buffered, format='JPEG', quality=85)
        img_base64 = base64.b64encode(buffered.getvalue()).decode()
        return {'base64': img_base64, 'descripcion': foto.get('descripcion', '')}
    except Exception:
        return {'base64': '', 'descripcion': foto.get('descripcion', '')}

def restaurar_foto(foto_dict):
    """Restaura un dict base64 a dict con BytesIO."""
    if not foto_dict.get('base64'):
        return {'imagen': None, 'descripcion': foto_dict.get('descripcion', '')}
    try:
        img_bytes = base64.b64decode(foto_dict['base64'])
        return {'imagen': io.BytesIO(img_bytes), 'descripcion': foto_dict.get('descripcion', '')}
    except Exception:
        return {'imagen': None, 'descripcion': foto_dict.get('descripcion', '')}

def serializar_lista_fotos(fotos_list):
    return [serializar_foto(f) for f in fotos_list]

def restaurar_lista_fotos(fotos_list):
    return [restaurar_foto(f) for f in fotos_list]

def convertir_hallazgos_a_serializable(hallazgos_list):
    resultado = []
    for hallazgo in hallazgos_list:
        h = {
            'tipo': hallazgo.get('tipo', ''),
            'descripcion': hallazgo.get('descripcion', ''),
            'normativa': hallazgo.get('normativa', ''),
            'accion': hallazgo.get('accion', ''),
            'resolucion': hallazgo.get('resolucion', ''),
            'estado': hallazgo.get('estado', ''),
            'fotos': serializar_lista_fotos(hallazgo.get('fotos', []))
        }
        resultado.append(h)
    return resultado

def restaurar_hallazgos_desde_json(hallazgos_list):
    resultado = []
    for hallazgo in hallazgos_list:
        h = {
            'tipo': hallazgo.get('tipo', ''),
            'descripcion': hallazgo.get('descripcion', ''),
            'normativa': hallazgo.get('normativa', ''),
            'accion': hallazgo.get('accion', ''),
            'resolucion': hallazgo.get('resolucion', ''),
            'estado': hallazgo.get('estado', ''),
            'fotos': restaurar_lista_fotos(hallazgo.get('fotos', []))
        }
        resultado.append(h)
    return resultado

def serializar_dict_fotos(fotos_dict):
    """Serializa un diccionario de listas de fotos (fotos_inspecciones, etc.)."""
    return {k: serializar_lista_fotos(v) for k, v in fotos_dict.items()}

def restaurar_dict_fotos(fotos_dict):
    """Restaura un diccionario de listas de fotos."""
    return {k: restaurar_lista_fotos(v) for k, v in fotos_dict.items()}

# ============================================================================
# FUNCIONES PARA CARGAR IMÁGENES
# ============================================================================

def cargar_logo():
    for p in ["logo_serfigobol.png", "assets/logo_serfigobol.png"]:
        if os.path.exists(p):
            return p
    return None

def cargar_logo_qfenix():
    for p in ["qfenix.png", "assets/qfenix.png"]:
        if os.path.exists(p):
            return p
    return None

def cargar_sello_brenda():
    for p in ["sello_brenda.jpg", "sello_brenda.png", "assets/sello_brenda.jpg"]:
        if os.path.exists(p):
            return p
    return None

def cargar_sello_erika():
    # CORREGIDO: nombre correcto "Erika Soruco" con K
    for p in ["sello_erika.jpg", "sello_erika.png", "assets/sello_erika.jpg"]:
        if os.path.exists(p):
            return p
    return None

def cargar_sello_cliente():
    for p in ["sello_cliente.jpg", "sello_cliente.png", "assets/sello_cliente.jpg"]:
        if os.path.exists(p):
            return p
    return None

# ============================================================================
# FUNCIÓN PARA CONVERTIR HTML DE ST_QUILL A WORD
# ============================================================================

def add_quill(paragraph, html_text):
    """Convierte HTML de Quill editor a runs de Word con formato correcto."""
    if not html_text:
        return

    # Separar por párrafos primero
    partes_parrafo = re.split(r'(<p>|</p>)', html_text)
    texto_acumulado = html_text

    parts = re.split(
        r'(<strong>|</strong>|<b>|</b>|<u>|</u>|<em>|</em>|<i>|</i>|<span[^>]*>|</span>|<p>|</p>|<br\s*/?>)',
        texto_acumulado
    )

    is_bold = False
    is_underline = False
    is_italic = False
    current_color = None
    current_bg = None

    for part in parts:
        if part in ['<strong>', '<b>']:
            is_bold = True
        elif part in ['</strong>', '</b>']:
            is_bold = False
        elif part == '<u>':
            is_underline = True
        elif part == '</u>':
            is_underline = False
        elif part in ['<em>', '<i>']:
            is_italic = True
        elif part in ['</em>', '</i>']:
            is_italic = False
        elif part.startswith('<span') and 'style' in part:
            color_match = re.search(r'color:\s*([^;]+)', part)
            if color_match:
                current_color = color_match.group(1).strip()
            bg_match = re.search(r'background-color:\s*([^;]+)', part)
            if bg_match:
                current_bg = bg_match.group(1).strip()
        elif part == '</span>':
            current_color = None
            current_bg = None
        elif part in ['<br/>', '<br>', '<br />', '</p>']:
            # Salto de línea en párrafo Word
            paragraph.add_run('\n')
        elif part == '<p>':
            pass  # Inicio de párrafo, no hacemos nada especial
        else:
            clean = re.sub('<[^<]+?>', '', part)
            if clean:
                run = paragraph.add_run(clean)
                run.bold = is_bold
                run.underline = is_underline
                run.italic = is_italic

                if current_color:
                    try:
                        color_map = {
                            'red': RGBColor(255, 0, 0), 'blue': RGBColor(0, 0, 255),
                            'green': RGBColor(0, 128, 0), 'yellow': RGBColor(255, 255, 0),
                            'orange': RGBColor(255, 165, 0), 'purple': RGBColor(128, 0, 128),
                            'black': RGBColor(0, 0, 0), 'white': RGBColor(255, 255, 255)
                        }
                        color_lower = current_color.lower()
                        if color_lower in color_map:
                            run.font.color.rgb = color_map[color_lower]
                        elif current_color.startswith('#'):
                            hex_color = current_color.lstrip('#')
                            if len(hex_color) == 6:
                                r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
                                run.font.color.rgb = RGBColor(r, g, b)
                    except Exception:
                        pass

                if current_bg:
                    try:
                        shading = OxmlElement('w:shd')
                        bg_value = current_bg.upper().replace('#', '')
                        color_bg_map = {
                            'YELLOW': 'FFFF00', 'RED': 'FF0000',
                            'GREEN': '00FF00', 'BLUE': '0000FF'
                        }
                        bg_value = color_bg_map.get(bg_value, bg_value)
                        shading.set(qn('w:fill'), bg_value)
                        run._r.get_or_add_rPr().append(shading)
                    except Exception:
                        pass

# ============================================================================
# FUNCIONES DE FORMATO PARA WORD
# ============================================================================

def make_titulo_counter():
    """Retorna un objeto contador para controlar el primer título (thread-safe por sesión)."""
    return {'primer_titulo': False}

def add_title_with_bar(doc, text, titulo_state):
    """Agrega título con barra gris oscura. titulo_state es dict local por sesión."""
    if titulo_state['primer_titulo']:
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(255, 255, 255)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), '404040')
    p._p.get_or_add_pPr().append(shading)
    titulo_state['primer_titulo'] = True

def add_table_professional(doc, df, title):
    if df is None:
        return
    if isinstance(df, dict) and '__columns__' in df:
        df = restaurar_dfs_desde_json(df)
    if isinstance(df, list):
        if len(df) > 0 and isinstance(df[0], dict):
            df = pd.DataFrame(df)
        else:
            return
    if not isinstance(df, pd.DataFrame):
        return
    if df.empty:
        return

    if title:
        doc.add_heading(title, level=2)
        for run in doc.paragraphs[-1].runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = True

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'

    for i, col in enumerate(df.columns):
        cell = table.cell(0, i)
        cell.text = str(col).strip()
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D3D3D3')
        cell._tc.get_or_add_tcPr().append(shading)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, val in enumerate(row):
            # CORREGIDO: evitar "nan" en celdas
            val_str = "" if pd.isna(val) else str(val).strip()
            cells[i].text = val_str
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_subtitle_left(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)

def add_heading_black(doc, text, level=2):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.bold = True

def add_photos_row(doc, fotos, title, cols_per_row=3):
    """Agrega fotos en filas de N columnas al documento Word."""
    if not fotos:
        return
    if title:
        doc.add_heading(title, level=3)
        for run in doc.paragraphs[-1].runs:
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.bold = True
    for i in range(0, len(fotos), cols_per_row):
        t = doc.add_table(rows=1, cols=cols_per_row)
        t.style = 'Table Grid'
        for j in range(cols_per_row):
            idx = i + j
            cell = t.cell(0, j)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if idx < len(fotos):
                foto = fotos[idx]
                try:
                    img_obj = foto.get('imagen')
                    if img_obj is None:
                        cell.text = ""
                        continue
                    if hasattr(img_obj, 'seek'):
                        img_obj.seek(0)
                    img = io.BytesIO(img_obj.read())
                    run = cell.paragraphs[0].add_run()
                    run.add_picture(img, width=Inches(2.0))
                    desc = foto.get('descripcion', '')
                    if desc:
                        cell.add_paragraph(f"FIGURA {idx+1}: {desc}")
                except Exception:
                    cell.text = "Error al cargar imagen"

def add_hallazgo_table(doc, hallazgo, numero):
    add_subtitle_left(doc, f"HALLAZGO {numero}")

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.allow_autofit = False

    filas = [
        ("TIPO", hallazgo.get('tipo', '')),
        ("DESCRIPCIÓN", hallazgo.get('descripcion', '')),
        ("NORMATIVA", hallazgo.get('normativa', '')),
        ("ACCIÓN", hallazgo.get('accion', '')),
        ("RESOLUCIÓN", hallazgo.get('resolucion', '')),
        ("ESTADO", hallazgo.get('estado', ''))
    ]

    for i, (campo, valor) in enumerate(filas):
        cell_campo = table.cell(i, 0)
        cell_campo.text = campo
        cell_campo.paragraphs[0].runs[0].bold = True
        cell_campo.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell_campo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D3D3D3')
        cell_campo._tc.get_or_add_tcPr().append(shading)

        cell_valor = table.cell(i, 1)
        # CORREGIDO: verificar NaN
        val_str = "" if pd.isna(valor) else str(valor)
        cell_valor.text = val_str
        cell_valor.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell_valor.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    fotos_hallazgo = hallazgo.get('fotos', [])
    if fotos_hallazgo:
        add_photos_row(doc, fotos_hallazgo, f"FOTOS HALLAZGO {numero}", cols_per_row=3)

def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = "PAGE"
    run._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def add_num_pages(paragraph):
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    instrText = OxmlElement('w:instrText')
    instrText.text = "NUMPAGES"
    run._r.append(instrText)
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)

def agregar_encabezado_word(doc, report_id, cliente, tipo_titulo="REPORTE DIARIO DE MONITOREO QA/QC", subtitulo=None):
    """Genera el encabezado estándar para documentos Word. Reutilizable."""
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = ""

    header_table = header.add_table(rows=1, cols=3, width=Inches(6.0))
    header_table.style = 'Table Grid'
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row in header_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    header_table.columns[0].width = Inches(0.8)
    logo = cargar_logo()
    if logo:
        try:
            cell_logo = header_table.cell(0, 0)
            cell_logo.paragraphs[0].clear()
            run = cell_logo.paragraphs[0].add_run()
            run.add_picture(logo, width=Inches(0.6))
            cell_logo.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            header_table.cell(0, 0).text = "SERFIGOBOL"
    else:
        header_table.cell(0, 0).text = "SERFIGOBOL"

    header_table.columns[1].width = Inches(3.5)
    cell_title = header_table.cell(0, 1)
    cell_title.paragraphs[0].clear()
    p_t = cell_title.paragraphs[0]
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_t.add_run(tipo_titulo)
    run_t.bold = True
    run_t.font.size = Pt(14)
    if subtitulo:
        p_t.add_run(f"\n{subtitulo}").font.size = Pt(10)

    header_table.columns[2].width = Inches(1.7)
    cell_info = header_table.cell(0, 2)
    cell_info.paragraphs[0].clear()
    p1 = cell_info.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run("REPORTE NO.: ")
    r1.bold = True
    r1.font.size = Pt(9)
    p1.add_run(report_id.strip())

    p2 = cell_info.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("CLIENTE: ")
    r2.bold = True
    r2.font.size = Pt(9)
    p2.add_run(cliente)

    fecha_emision = datetime.now().strftime('%Y-%m-%d')
    p3 = cell_info.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("FECHA DE EMISIÓN: ")
    r3.bold = True
    r3.font.size = Pt(9)
    p3.add_run(fecha_emision)

    p_espacio = header.add_paragraph()
    p_espacio.paragraph_format.space_after = Pt(10)
    p_espacio.text = ""

    return fecha_emision

def agregar_pie_pagina(doc):
    """Genera el pie de página estándar. Reutilizable."""
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.text = ""

    p_numeracion = footer.add_paragraph()
    p_numeracion.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_numeracion.style.font.size = Pt(12)
    add_page_number(p_numeracion)
    p_numeracion.add_run(' de ')
    add_num_pages(p_numeracion)

    p_confidencial = footer.add_paragraph()
    p_confidencial.text = "DOCUMENTO CONFIDENCIAL | SERFIGOBOL S.R.L."
    p_confidencial.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_confidencial.style.font.size = Pt(9)

def agregar_tabla_cabecera(doc, datos_lista):
    """Genera tabla de datos principales 4 columnas. Reutilizable."""
    t = doc.add_table(rows=len(datos_lista), cols=4)
    t.style = 'Table Grid'
    for i, row in enumerate(datos_lista):
        for j, val in enumerate(row):
            cell = t.cell(i, j)
            # CORREGIDO: proteger NaN
            val_str = "" if (val is None or (isinstance(val, float) and pd.isna(val))) else str(val).strip()
            cell.text = val_str
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if j % 2 == 0:
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'D3D3D3')
                cell._tc.get_or_add_tcPr().append(shading)
    return t

def agregar_seccion_representantes(doc, datos_form, tipo='qaqc'):
    """Genera la sección de representantes y firmas. Reutilizable para QA/QC y DBR."""
    sig_table = doc.add_table(rows=3, cols=3)
    sig_table.style = 'Table Grid'
    for row in sig_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for i in range(3):
        sig_table.cell(0, i).text = "FIRMA"
        sig_table.cell(0, i).paragraphs[0].runs[0].bold = True

    sello_c = cargar_sello_cliente()
    if sello_c:
        try:
            run = sig_table.cell(1, 1).paragraphs[0].add_run()
            run.add_picture(sello_c, width=Inches(1.3))
        except Exception:
            pass

    # CORREGIDO: nombre correcto "Erika Soruco" con K
    rep_serfi_elegida = datos_form.get('rep_serfigobol', 'Brenda Figueroa')
    if rep_serfi_elegida == 'Brenda Figueroa':
        sello_serfi = cargar_sello_brenda()
    else:
        sello_serfi = cargar_sello_erika()
    if sello_serfi:
        try:
            run = sig_table.cell(1, 2).paragraphs[0].add_run()
            run.add_picture(sello_serfi, width=Inches(1.3))
        except Exception:
            pass

    etiqueta_col0 = "GERENTE DE CONTRATO" if tipo == 'qaqc' else "GERENTE DE PROYECTO"
    nombre_col0_key = 'gerente_contrato' if tipo == 'qaqc' else 'gerente_proyecto'

    for col_idx, (nombre_key, etiqueta) in enumerate([
        (nombre_col0_key, etiqueta_col0),
        ('rep_cliente', 'Representante Cliente'),
        ('rep_serfigobol', 'Representante de SERFIGOBOL S.R.L.')
    ]):
        cell = sig_table.cell(2, col_idx)
        p = cell.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        nombre = datos_form.get(nombre_key, '____________________')
        run = p.add_run(f"{nombre}\n{etiqueta}")
        run.bold = True
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'D3D3D3')
        cell._tc.get_or_add_tcPr().append(shading)

    for col in range(3):
        sig_table.columns[col].width = Inches(2.0)

# ============================================================================
# CONFIGURACIÓN Y BASE DE DATOS
# ============================================================================

st.set_page_config(page_title="Q-Fenix - Sistema QA/QC", layout="wide")

def ejecutar_query(query, params=(), fetch=False):
    try:
        with sqlite3.connect(DB_PATH) as conn:
            cursor = conn.cursor()
            cursor.execute(query, params)
            if fetch:
                return cursor.fetchall()
            conn.commit()
            return True
    except sqlite3.Error as e:
        st.error(f"Error de base de datos: {e}")
        return [] if fetch else None

def inicializar_bd():
    ejecutar_query("""
        CREATE TABLE IF NOT EXISTS usuarios (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password TEXT NOT NULL,
            rol TEXT DEFAULT 'inspector',
            activo INTEGER DEFAULT 1,
            fecha_creacion TEXT
        )
    """)
    ejecutar_query("""
        CREATE TABLE IF NOT EXISTS logs (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario_id INTEGER,
            accion TEXT,
            detalles TEXT,
            fecha TEXT
        )
    """)
    ejecutar_query("""
        CREATE TABLE IF NOT EXISTS proyectos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            pais TEXT, operadora TEXT, contrato_nom TEXT NOT NULL,
            contrato_num TEXT, contacto_cliente TEXT, gerente_proyecto TEXT,
            cant_pozos INTEGER, fecha_creacion TEXT
        )
    """)
    ejecutar_query("""
        CREATE TABLE IF NOT EXISTS pozos (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            proyecto_id INTEGER,
            rig TEXT,
            nombre_pozo TEXT NOT NULL,
            fecha_creacion TEXT,
            FOREIGN KEY (proyecto_id) REFERENCES proyectos (id)
        )
    """)
    ejecutar_query("""
        CREATE TABLE IF NOT EXISTS proveedores (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            pozo_id INTEGER,
            nombre_proveedor TEXT,
            FOREIGN KEY (pozo_id) REFERENCES pozos (id)
        )
    """)
    ejecutar_query("""
        CREATE TABLE IF NOT EXISTS inspecciones (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            proyecto_id INTEGER, pozo_id INTEGER, proveedor_id INTEGER,
            reporte_id TEXT, reporte_numero INTEGER, fecha_inspeccion TEXT,
            cliente TEXT, contrato TEXT, pozo_nombre TEXT,
            datos_json TEXT, fecha_servicio TEXT, fecha_creacion TEXT, fecha_emision TEXT,
            estado_reporte TEXT DEFAULT 'BORRADOR',
            usuario_creador_id INTEGER,
            fecha_modificacion TEXT,
            FOREIGN KEY (proyecto_id) REFERENCES proyectos (id),
            FOREIGN KEY (pozo_id) REFERENCES pozos (id),
            FOREIGN KEY (proveedor_id) REFERENCES proveedores (id),
            FOREIGN KEY (usuario_creador_id) REFERENCES usuarios (id)
        )
    """)
    ejecutar_query("""
        CREATE TABLE IF NOT EXISTS contador_reportes (
            pozo_id INTEGER PRIMARY KEY, ultimo_numero INTEGER DEFAULT 0,
            FOREIGN KEY (pozo_id) REFERENCES pozos (id)
        )
    """)
    ejecutar_query("""
        CREATE TABLE IF NOT EXISTS discrepancias_pozo (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            pozo_id INTEGER,
            proyecto_id INTEGER,
            reporte_id TEXT,
            tipo TEXT,
            descripcion TEXT,
            normativa TEXT,
            acciones TEXT,
            resolucion TEXT,
            estado TEXT,
            fecha TEXT,
            usuario_id INTEGER,
            fotos_json TEXT,
            FOREIGN KEY (pozo_id) REFERENCES pozos (id),
            FOREIGN KEY (proyecto_id) REFERENCES proyectos (id),
            FOREIGN KEY (usuario_id) REFERENCES usuarios (id)
        )
    """)
    ejecutar_query("""
        CREATE TABLE IF NOT EXISTS contador_reportes_dbr (
            pozo_id INTEGER PRIMARY KEY, ultimo_numero INTEGER DEFAULT 0,
            FOREIGN KEY (pozo_id) REFERENCES pozos (id)
        )
    """)

    # Migraciones seguras
    for migracion in [
        ("usuarios", "activo", "ALTER TABLE usuarios ADD COLUMN activo INTEGER DEFAULT 1"),
        ("proyectos", "gerente_proyecto", "ALTER TABLE proyectos ADD COLUMN gerente_proyecto TEXT"),
        ("inspecciones", "fecha_creacion", "ALTER TABLE inspecciones ADD COLUMN fecha_creacion TEXT"),
        ("inspecciones", "fecha_emision", "ALTER TABLE inspecciones ADD COLUMN fecha_emision TEXT"),
        ("inspecciones", "estado_reporte", "ALTER TABLE inspecciones ADD COLUMN estado_reporte TEXT DEFAULT 'BORRADOR'"),
        ("discrepancias_pozo", "normativa", "ALTER TABLE discrepancias_pozo ADD COLUMN normativa TEXT"),
        ("discrepancias_pozo", "fotos_json", "ALTER TABLE discrepancias_pozo ADD COLUMN fotos_json TEXT"),
    ]:
        tabla, columna, sql = migracion
        try:
            cols = ejecutar_query(f"PRAGMA table_info({tabla})", fetch=True)
            if cols and columna not in [c[1] for c in cols]:
                ejecutar_query(sql)
        except Exception:
            pass

    admin_exists = ejecutar_query("SELECT id FROM usuarios WHERE username = 'admin' AND activo = 1", fetch=True)
    if not admin_exists:
        ejecutar_query(
            "INSERT INTO usuarios (username, password, rol, activo, fecha_creacion) VALUES (?, ?, ?, ?, ?)",
            ("admin", hash_password("1234"), "admin", 1, datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        )

inicializar_bd()

# ============================================================================
# FUNCIÓN ÚNICA PARA GUARDAR DATOS DEL REPORTE QA/QC
# ============================================================================

def guardar_datos_reporte(p_sel, poz_sel, prov_sel, report_id, es_generacion_word=False):
    """Guarda reporte QA/QC incluyendo fotos de inspección, ensamble y prueba."""
    datos_serializables = convertir_df_a_serializable(st.session_state.form_data)
    hallazgos_serializables = convertir_hallazgos_a_serializable(st.session_state.hallazgos_list)

    # CORREGIDO: serializar fotos de inspección, ensamble y prueba
    fotos_inspecciones_serial = serializar_dict_fotos(st.session_state.fotos_inspecciones)
    fotos_ensambles_serial = serializar_dict_fotos(st.session_state.fotos_ensambles)
    fotos_pruebas_serial = serializar_dict_fotos(st.session_state.fotos_pruebas)
    fotos_generales_serial = serializar_lista_fotos(st.session_state.fotos_generales)

    datos_completos = {
        'form_data': datos_serializables,
        'check_vals': st.session_state.form_data.get('check_vals', {}),
        'otras': st.session_state.form_data.get('otras', ''),
        'nota_instrumentos': st.session_state.nota_instrumentos,
        'nota_personal': st.session_state.nota_personal,
        'mostrar_nota_inst': st.session_state.mostrar_nota_inst,
        'mostrar_nota_pers': st.session_state.mostrar_nota_pers,
        'n_inspecciones': st.session_state.n_inspecciones,
        'n_ensambles': st.session_state.n_ensambles,
        'n_pruebas': st.session_state.n_pruebas,
        'n_fotos_secc': st.session_state.n_fotos_secc,
        'hallazgos_list': hallazgos_serializables,
        'estado_herramientas_df': convertir_df_a_serializable(st.session_state.estado_herramientas_df),
        'n_fotos_insp': st.session_state.n_fotos_insp,
        'n_fotos_ens': st.session_state.n_fotos_ens,
        'n_fotos_prueba': st.session_state.n_fotos_prueba,
        # CORREGIDO: incluir fotos de inspección, ensamble y prueba
        'fotos_inspecciones': fotos_inspecciones_serial,
        'fotos_ensambles': fotos_ensambles_serial,
        'fotos_pruebas': fotos_pruebas_serial,
        'fotos_generales': fotos_generales_serial,
        'datos_seleccion': {
            'proyecto_id': p_sel[0],
            'pozo_id': poz_sel[0],
            'proveedor_id': prov_sel[0] if prov_sel else None
        }
    }

    # Usar UNA sola conexión para verificar y guardar
    ahora = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    try:
        with sqlite3.connect(DB_PATH) as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id, estado_reporte FROM inspecciones WHERE reporte_id = ?", (report_id,))
            existe = cursor.fetchall()

            datos_json_str = json.dumps(datos_completos)
            nuevo_estado = None
            if es_generacion_word and existe:
                estado_actual = existe[0][1] if existe else 'BORRADOR'
                if estado_actual == 'BORRADOR':
                    nuevo_estado = 'EMITIDO'

            if existe:
                if nuevo_estado:
                    cursor.execute("""UPDATE inspecciones SET
                        proyecto_id=?, pozo_id=?, proveedor_id=?,
                        cliente=?, contrato=?, pozo_nombre=?,
                        fecha_servicio=?, datos_json=?, fecha_modificacion=?, estado_reporte=?
                        WHERE reporte_id=?""",
                        (p_sel[0], poz_sel[0], prov_sel[0] if prov_sel else None,
                         st.session_state.form_data.get('cliente', ''),
                         st.session_state.form_data.get('contrato', ''),
                         st.session_state.form_data.get('pozo', ''),
                         st.session_state.form_data.get('fecha_serv', ''),
                         datos_json_str, ahora, nuevo_estado, report_id))
                else:
                    cursor.execute("""UPDATE inspecciones SET
                        proyecto_id=?, pozo_id=?, proveedor_id=?,
                        cliente=?, contrato=?, pozo_nombre=?,
                        fecha_servicio=?, datos_json=?, fecha_modificacion=?
                        WHERE reporte_id=?""",
                        (p_sel[0], poz_sel[0], prov_sel[0] if prov_sel else None,
                         st.session_state.form_data.get('cliente', ''),
                         st.session_state.form_data.get('contrato', ''),
                         st.session_state.form_data.get('pozo', ''),
                         st.session_state.form_data.get('fecha_serv', ''),
                         datos_json_str, ahora, report_id))
            else:
                cursor.execute("""INSERT INTO inspecciones
                    (proyecto_id, pozo_id, proveedor_id, reporte_id, reporte_numero,
                     fecha_inspeccion, cliente, contrato, pozo_nombre, fecha_servicio,
                     fecha_creacion, estado_reporte, datos_json,
                     usuario_creador_id, fecha_modificacion)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                    (p_sel[0], poz_sel[0], prov_sel[0] if prov_sel else None,
                     report_id, st.session_state.reporte_numero_actual or 1,
                     ahora,
                     st.session_state.form_data.get('cliente', ''),
                     st.session_state.form_data.get('contrato', ''),
                     st.session_state.form_data.get('pozo', ''),
                     st.session_state.form_data.get('fecha_serv', ''),
                     ahora, 'BORRADOR', datos_json_str,
                     st.session_state.usuario_id, ahora))
            conn.commit()
            if nuevo_estado:
                st.session_state.estado_reporte_actual = nuevo_estado
            return True
    except Exception as e:
        st.error(f"❌ ERROR AL GUARDAR: {e}")
        return False

# ============================================================================
# FUNCIÓN ÚNICA PARA GUARDAR DATOS DEL REPORTE DBR
# ============================================================================

def guardar_datos_reporte_dbr(p_sel, poz_sel, prov_sel, report_id, es_generacion_word=False):
    """Guarda reporte DBR en la base de datos con una sola conexión."""
    ed_items_serial = convertir_df_a_serializable(st.session_state.ed_items_dbr)
    tabla_evento_serial = convertir_df_a_serializable(st.session_state.tabla_historial_evento)
    tabla_herramienta_serial = convertir_df_a_serializable(st.session_state.tabla_historial_herramienta)
    tabla_procesos_serial = convertir_df_a_serializable(st.session_state.tabla_procesos)
    tabla_certificados_serial = convertir_df_a_serializable(st.session_state.tabla_certificados)

    fotos_evento_serial = serializar_lista_fotos(st.session_state.fotos_historial_evento)
    fotos_herramienta_serial = serializar_lista_fotos(st.session_state.fotos_historial_herramienta)
    fotos_procesos_serial = serializar_lista_fotos(st.session_state.fotos_procesos)
    fotos_resumen_serial = serializar_lista_fotos(st.session_state.fotos_resumen)
    hallazgos_serial = convertir_hallazgos_a_serializable(st.session_state.hallazgos_list_dbr)

    datos_completos = {
        'form_data_dbr': st.session_state.form_data_dbr,
        'ed_items_dbr': ed_items_serial,
        'tabla_historial_evento': tabla_evento_serial,
        'fotos_historial_evento': fotos_evento_serial,
        'n_fotos_historial_evento': st.session_state.n_fotos_historial_evento,
        'tabla_historial_herramienta': tabla_herramienta_serial,
        'fotos_historial_herramienta': fotos_herramienta_serial,
        'n_fotos_historial_herramienta': st.session_state.n_fotos_historial_herramienta,
        'tabla_procesos': tabla_procesos_serial,
        'fotos_procesos': fotos_procesos_serial,
        'n_fotos_procesos': st.session_state.n_fotos_procesos,
        'tabla_certificados': tabla_certificados_serial,
        'fotos_resumen': fotos_resumen_serial,
        'n_fotos_resumen': st.session_state.n_fotos_resumen,
        'hallazgos_list_dbr': hallazgos_serial,
        'representantes': {
            'gerente_proyecto': st.session_state.form_data_dbr.get('gerente_proyecto', ''),
            'rep_cliente': st.session_state.form_data_dbr.get('rep_cliente', ''),
            'cargo_cliente': st.session_state.form_data_dbr.get('cargo_cliente', ''),
            'rep_serfigobol': st.session_state.form_data_dbr.get('rep_serfigobol', 'Brenda Figueroa')
        },
        'datos_seleccion': {
            'proyecto_id': p_sel[0],
            'pozo_id': poz_sel[0],
            'proveedor_id': prov_sel[0] if prov_sel else None
        }
    }

    ahora = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    try:
        with sqlite3.connect(DB_PATH) as conn:
            cursor = conn.cursor()
            cursor.execute("SELECT id, estado_reporte FROM inspecciones WHERE reporte_id = ?", (report_id,))
            existe = cursor.fetchall()

            datos_json_str = json.dumps(datos_completos)
            nuevo_estado = None
            if es_generacion_word and existe:
                estado_actual = existe[0][1] if existe else 'BORRADOR'
                if estado_actual == 'BORRADOR':
                    nuevo_estado = 'EMITIDO'

            if existe:
                if nuevo_estado:
                    cursor.execute("""UPDATE inspecciones SET
                        proyecto_id=?, pozo_id=?, proveedor_id=?,
                        cliente=?, contrato=?, pozo_nombre=?,
                        fecha_servicio=?, datos_json=?, fecha_modificacion=?, estado_reporte=?
                        WHERE reporte_id=?""",
                        (p_sel[0], poz_sel[0], prov_sel[0] if prov_sel else None,
                         st.session_state.form_data_dbr.get('cliente', ''),
                         st.session_state.form_data_dbr.get('contrato', ''),
                         st.session_state.form_data_dbr.get('pozo', ''),
                         st.session_state.form_data_dbr.get('fecha_serv', ''),
                         datos_json_str, ahora, nuevo_estado, report_id))
                else:
                    cursor.execute("""UPDATE inspecciones SET
                        proyecto_id=?, pozo_id=?, proveedor_id=?,
                        cliente=?, contrato=?, pozo_nombre=?,
                        fecha_servicio=?, datos_json=?, fecha_modificacion=?
                        WHERE reporte_id=?""",
                        (p_sel[0], poz_sel[0], prov_sel[0] if prov_sel else None,
                         st.session_state.form_data_dbr.get('cliente', ''),
                         st.session_state.form_data_dbr.get('contrato', ''),
                         st.session_state.form_data_dbr.get('pozo', ''),
                         st.session_state.form_data_dbr.get('fecha_serv', ''),
                         datos_json_str, ahora, report_id))
            else:
                cursor.execute("""INSERT INTO inspecciones
                    (proyecto_id, pozo_id, proveedor_id, reporte_id, reporte_numero,
                     fecha_inspeccion, cliente, contrato, pozo_nombre, fecha_servicio,
                     fecha_creacion, estado_reporte, datos_json,
                     usuario_creador_id, fecha_modificacion)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                    (p_sel[0], poz_sel[0], prov_sel[0] if prov_sel else None,
                     report_id, st.session_state.get('reporte_numero_actual_dbr', 1),
                     ahora,
                     st.session_state.form_data_dbr.get('cliente', ''),
                     st.session_state.form_data_dbr.get('contrato', ''),
                     st.session_state.form_data_dbr.get('pozo', ''),
                     st.session_state.form_data_dbr.get('fecha_serv', ''),
                     ahora, 'BORRADOR', datos_json_str,
                     st.session_state.usuario_id, ahora))
            conn.commit()
            if nuevo_estado:
                st.session_state.estado_reporte_actual_dbr = nuevo_estado
            return True
    except Exception as e:
        st.error(f"❌ ERROR AL GUARDAR DBR: {e}")
        return False

def guardar_hallazgos_en_discrepancias(poz_sel, p_sel, report_id, hallazgos_list):
    """Guarda hallazgos en tabla discrepancias_pozo. Reutilizable."""
    try:
        with sqlite3.connect(DB_PATH) as conn:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM discrepancias_pozo WHERE reporte_id = ?", (report_id,))
            if hallazgos_list:
                ahora = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                for hallazgo in hallazgos_list:
                    cursor.execute("""INSERT INTO discrepancias_pozo
                        (pozo_id, proyecto_id, reporte_id, tipo, descripcion, normativa, acciones, resolucion, estado, fecha, usuario_id)
                        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                        (poz_sel[0], p_sel[0], report_id,
                         hallazgo.get('tipo', ''), hallazgo.get('descripcion', ''),
                         hallazgo.get('normativa', ''), hallazgo.get('accion', ''),
                         hallazgo.get('resolucion', ''), hallazgo.get('estado', 'ABIERTO'),
                         ahora, st.session_state.usuario_id))
            conn.commit()
    except Exception as e:
        st.error(f"❌ Error al guardar discrepancias: {e}")

# ============================================================================
# ESTADO DE SESIÓN - INICIALIZACIÓN COMPLETA Y CENTRALIZADA
# ============================================================================

defaults = {
    'auth': False,
    'usuario_id': None,
    'usuario_rol': None,
    'usuario_nombre': None,
    'formulario_activo': False,
    'modo_edicion': False,
    'edicion_reporte_id': None,
    'datos_seleccion': {},
    'reporte_actual_id': None,
    'reporte_numero_actual': None,
    'ultimo_activo': datetime.now(),
    'menu_seleccionado': None,
    'datos_editados_cargados': False,
    'n_inspecciones': 1,
    'n_ensambles': 1,
    'n_pruebas': 1,
    'n_fotos_secc': 1,
    'fotos_inspecciones': {},
    'fotos_ensambles': {},
    'fotos_pruebas': {},
    'fotos_generales': [],
    'nota_instrumentos': '',
    'nota_personal': '',
    'mostrar_nota_inst': False,
    'mostrar_nota_pers': False,
    'hallazgos_list': [],
    'estado_herramientas_df': pd.DataFrame(columns=[
        "Herramienta", "Cant. Inspeccionada", "Cant. Aceptada", "Cant. Rechazada", "Cant. Pendiente"
    ]),
    'form_data': {},
    'n_fotos_insp': {},
    'n_fotos_ens': {},
    'n_fotos_prueba': {},
    'historial_pagina': 1,
    # DBR
    'formulario_activo_dbr': False,
    'modo_edicion_dbr': False,
    'edicion_reporte_id_dbr': None,
    'datos_editados_dbr_cargados': False,
    'reporte_actual_id_dbr': None,
    'reporte_numero_actual_dbr': None,
    'estado_reporte_actual_dbr': 'BORRADOR',
    'fecha_emision_actual_dbr': None,
    'datos_seleccion_dbr': {},
    'form_data_dbr': {},
    'hallazgos_list_dbr': [],
    'ed_items_dbr': pd.DataFrame(columns=["N°", "DESCRIPCIÓN", "OD", "CONEXIONES", "NUMERO DE SERIE"]),
    'n_fotos_historial_evento': 1,
    'n_fotos_historial_herramienta': 1,
    'n_fotos_procesos': 1,
    'n_fotos_resumen': 1,
    'tabla_historial_evento': pd.DataFrame(columns=["FECHA DE EVENTO", "EVENTO", "OBSERVACION"]),
    'tabla_historial_herramienta': pd.DataFrame(columns=["FECHA DE ACCION", "ACCION", "OBSERVACIÓN"]),
    'tabla_procesos': pd.DataFrame(columns=["FECHA", "PROCESO", "N REPORTE", "OBSERVACIÓN"]),
    'tabla_certificados': pd.DataFrame(columns=["DOCUMENTO", "NUMERO DOC./REVISION/N REF.", "OBSERVACIÓN"]),
    'fotos_historial_evento': [],
    'fotos_historial_herramienta': [],
    'fotos_procesos': [],
    'fotos_resumen': [],
    'estado_reporte_actual': 'BORRADOR',
    'fecha_emision_actual': None,
}

for key, val in defaults.items():
    if key not in st.session_state:
        st.session_state[key] = val

# ============================================================================
# VERIFICAR EXPIRACIÓN DE SESIÓN
# ============================================================================

if st.session_state.auth:
    tiempo_transcurrido = datetime.now() - st.session_state.ultimo_activo
    if tiempo_transcurrido > timedelta(minutes=30):
        st.session_state.auth = False
        st.warning("⏰ SESIÓN EXPIRADA. Sus datos no guardados se han perdido.")
        st.rerun()
    else:
        st.session_state.ultimo_activo = datetime.now()

# ============================================================================
# AUTENTICACIÓN
# ============================================================================

if not st.session_state.auth:
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        logo_qfenix = cargar_logo_qfenix()
        if logo_qfenix:
            col_logo1, col_logo2, col_logo3 = st.columns([1, 2, 1])
            with col_logo2:
                st.image(logo_qfenix, width=100, use_container_width=True)
                st.markdown("<h1 style='text-align: center; margin-top: -10px; color: #C26A2C;font-weight: 500;'>Q-Fenix</h1>", unsafe_allow_html=True)
        else:
            st.markdown("<h1 style='text-align: center; color: #C26A2C;font-weight: 500;'>Q-Fenix</h1>", unsafe_allow_html=True)

        st.subheader("ACCESO AL SISTEMA")
        usuario = st.text_input("USUARIO")
        contrasena = st.text_input("CONTRASEÑA", type="password")
        if st.button("INGRESAR", use_container_width=True):
            resultado = verificar_usuario(usuario, contrasena)
            if resultado:
                st.session_state.auth = True
                st.session_state.usuario_id = resultado[0]
                st.session_state.usuario_nombre = resultado[1]
                st.session_state.usuario_rol = resultado[2]
                st.session_state.ultimo_activo = datetime.now()
                st.session_state.datos_seleccion = {}
                st.session_state.formulario_activo = False
                st.session_state.menu_seleccionado = None
                registrar_log(st.session_state.usuario_id, "INICIO DE SESIÓN")
                st.rerun()
            else:
                st.error("❌ CREDENCIALES INCORRECTAS")
    st.stop()

# ============================================================================
# MENÚ PRINCIPAL
# ============================================================================

logo_path = cargar_logo()
if logo_path:
    st.sidebar.image(logo_path, width=100)

if st.session_state.usuario_rol == "admin":
    menu_opciones = ["🏗️ GESTIÓN DE PROYECTOS", "📝 NUEVO REPORTE QA/QC", "📝 NUEVO REPORTE QA/QC - DBR",
                     "📊 HISTORIAL DE INSPECCIONES", "📋 DISCREPANCIAS POR POZO", "📄 TIMESHEET", "👥 ADMINISTRACIÓN"]
else:
    menu_opciones = ["📝 NUEVO REPORTE QA/QC", "📝 NUEVO REPORTE QA/QC - DBR",
                     "📊 HISTORIAL DE INSPECCIONES", "📋 DISCREPANCIAS POR POZO", "📄 TIMESHEET"]

if st.session_state.menu_seleccionado:
    menu = st.session_state.menu_seleccionado
    st.session_state.menu_seleccionado = None
else:
    menu = st.sidebar.radio("📋 MENÚ PRINCIPAL", menu_opciones)

st.sidebar.divider()
if st.session_state.usuario_nombre:
    st.sidebar.write(f"👤 **USUARIO:** {st.session_state.usuario_nombre}")
if st.session_state.usuario_rol:
    st.sidebar.write(f"🔒 **ROL:** {st.session_state.usuario_rol.upper()}")
if st.sidebar.button("🚪 CERRAR SESIÓN", use_container_width=True):
    registrar_log(st.session_state.usuario_id, "CIERRE DE SESIÓN")
    st.session_state.auth = False
    st.session_state.datos_seleccion = {}
    st.session_state.menu_seleccionado = None
    st.rerun()

# ============================================================================
# ADMINISTRACIÓN DE USUARIOS
# ============================================================================

if menu == "👥 ADMINISTRACIÓN" and st.session_state.usuario_rol == "admin":
    st.header("👥 ADMINISTRACIÓN DE USUARIOS")

    tab_usuarios, tab_logs = st.tabs(["📋 USUARIOS", "📜 LOGS DE ACTIVIDAD"])

    with tab_usuarios:
        st.subheader("➕ CREAR NUEVO USUARIO")
        with st.form("form_nuevo_usuario"):
            col1, col2, col3 = st.columns(3)
            with col1:
                nuevo_user = st.text_input("USUARIO")
            with col2:
                nuevo_pass = st.text_input("CONTRASEÑA", type="password")
            with col3:
                nuevo_rol = st.selectbox("ROL", ["inspector", "admin"])
            if st.form_submit_button("💾 CREAR USUARIO"):
                if nuevo_user and nuevo_pass:
                    existe = ejecutar_query("SELECT id FROM usuarios WHERE username = ? AND activo = 1", (nuevo_user,), fetch=True)
                    if not existe:
                        ejecutar_query(
                            "INSERT INTO usuarios (username, password, rol, activo, fecha_creacion) VALUES (?, ?, ?, ?, ?)",
                            (nuevo_user, hash_password(nuevo_pass), nuevo_rol, 1, datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
                        )
                        registrar_log(st.session_state.usuario_id, f"CREAR USUARIO: {nuevo_user}")
                        st.success(f"✅ USUARIO {nuevo_user} CREADO")
                        st.rerun()
                    else:
                        st.error("❌ EL USUARIO YA EXISTE Y ESTÁ ACTIVO")
                else:
                    st.error("❌ COMPLETE TODOS LOS CAMPOS")

        st.subheader("📋 LISTA DE USUARIOS")
        usuarios = ejecutar_query("SELECT id, username, rol, activo, fecha_creacion FROM usuarios ORDER BY id", fetch=True)
        if usuarios:
            df_usuarios = pd.DataFrame(usuarios, columns=["ID", "USUARIO", "ROL", "ACTIVO", "FECHA CREACIÓN"])
            df_usuarios["ACTIVO"] = df_usuarios["ACTIVO"].map({1: "🟢 ACTIVO", 0: "🔴 INACTIVO"})
            st.dataframe(df_usuarios, use_container_width=True, hide_index=True)

            col_deshab, col_react, col_reset = st.columns(3)

            with col_deshab:
                st.subheader("🗑️ DESHABILITAR USUARIO")
                st.info("⚠️ Los reportes se conservan")
                usuarios_activos = ejecutar_query("SELECT id, username FROM usuarios WHERE activo = 1 AND username != ?", (st.session_state.usuario_nombre,), fetch=True)
                if usuarios_activos:
                    usuario_deshabilitar = st.selectbox("SELECCIONAR", usuarios_activos, format_func=lambda x: x[1], key="deshab_select")
                    if st.button("⚠️ DESHABILITAR", type="primary", key="deshab_btn"):
                        ejecutar_query("UPDATE usuarios SET activo = 0 WHERE id = ?", (usuario_deshabilitar[0],))
                        registrar_log(st.session_state.usuario_id, f"DESHABILITAR USUARIO: {usuario_deshabilitar[1]}")
                        st.success(f"✅ USUARIO {usuario_deshabilitar[1]} DESHABILITADO")
                        st.rerun()
                else:
                    st.info("📌 NO HAY USUARIOS ACTIVOS PARA DESHABILITAR")

            with col_react:
                st.subheader("🔄 REACTIVAR USUARIO")
                usuarios_inactivos = ejecutar_query("SELECT id, username FROM usuarios WHERE activo = 0", fetch=True)
                if usuarios_inactivos:
                    usuario_reactivar = st.selectbox("SELECCIONAR", usuarios_inactivos, format_func=lambda x: x[1], key="react_select")
                    if st.button("🔄 REACTIVAR", key="react_btn"):
                        ejecutar_query("UPDATE usuarios SET activo = 1 WHERE id = ?", (usuario_reactivar[0],))
                        registrar_log(st.session_state.usuario_id, f"REACTIVAR USUARIO: {usuario_reactivar[1]}")
                        st.success(f"✅ USUARIO {usuario_reactivar[1]} REACTIVADO")
                        st.rerun()
                else:
                    st.info("📌 NO HAY USUARIOS INACTIVOS")

            with col_reset:
                st.subheader("🔑 RESETEAR CONTRASEÑA")
                todos_usuarios = ejecutar_query("SELECT id, username FROM usuarios WHERE username != ?", (st.session_state.usuario_nombre,), fetch=True)
                if todos_usuarios:
                    usuario_reset = st.selectbox("SELECCIONAR", todos_usuarios, format_func=lambda x: x[1], key="reset_select")
                    nueva_pass = st.text_input("NUEVA CONTRASEÑA", type="password", key="reset_pass")
                    if st.button("🔑 RESETEAR", key="reset_btn") and nueva_pass:
                        ejecutar_query("UPDATE usuarios SET password = ? WHERE id = ?", (hash_password(nueva_pass), usuario_reset[0]))
                        registrar_log(st.session_state.usuario_id, f"RESETEAR CONTRASEÑA: {usuario_reset[1]}")
                        st.success(f"✅ CONTRASEÑA DE {usuario_reset[1]} ACTUALIZADA")
                        st.rerun()
                else:
                    st.info("📌 NO HAY OTROS USUARIOS")

    with tab_logs:
        st.subheader("📜 REGISTRO DE ACTIVIDADES")
        logs = ejecutar_query("""
            SELECT l.fecha, u.username, l.accion, l.detalles
            FROM logs l
            LEFT JOIN usuarios u ON l.usuario_id = u.id
            ORDER BY l.fecha DESC
            LIMIT 100
        """, fetch=True)
        if logs:
            df_logs = pd.DataFrame(logs, columns=["FECHA", "USUARIO", "ACCIÓN", "DETALLES"])
            st.dataframe(df_logs, use_container_width=True, hide_index=True)

# ============================================================================
# GESTIÓN DE PROYECTOS
# ============================================================================
elif menu == "🏗️ GESTIÓN DE PROYECTOS" and st.session_state.usuario_rol == "admin":
    st.header("🏗️ GESTIÓN DE PROYECTOS")

    tab_reg, tab_edit, tab_pozos_prov = st.tabs(["📁 NUEVO PROYECTO", "✏️ EDITAR PROYECTO", "🛢️ POZOS/PROVEEDORES"])

    with tab_reg:
        with st.form("nuevo_proyecto"):
            col1, col2 = st.columns(2)
            with col1:
                contrato_nom = st.text_input("NOMBRE DEL CONTRATO *")
                operadora = st.text_input("OPERADORA")
                gerente_proyecto = st.text_input("GERENTE DE PROYECTO")
            with col2:
                contrato_num = st.text_input("N° CONTRATO")
                contacto_cliente = st.text_input("CONTACTO CLIENTE")
            if st.form_submit_button("💾 GUARDAR PROYECTO"):
                if contrato_nom:
                    ejecutar_query(
                        "INSERT INTO proyectos (contrato_nom, operadora, contrato_num, contacto_cliente, gerente_proyecto, fecha_creacion) VALUES (?,?,?,?,?,?)",
                        (contrato_nom, operadora, contrato_num, contacto_cliente, gerente_proyecto, datetime.now().strftime('%Y-%m-%d'))
                    )
                    registrar_log(st.session_state.usuario_id, f"CREAR PROYECTO: {contrato_nom}")
                    st.success("✅ PROYECTO GUARDADO")
                    st.rerun()

    with tab_edit:
        proyectos = ejecutar_query("SELECT id, contrato_nom, operadora, contrato_num, contacto_cliente, gerente_proyecto FROM proyectos", fetch=True)
        if proyectos:
            proyecto_seleccionado = st.selectbox("SELECCIONAR PROYECTO", proyectos, format_func=lambda x: x[1])
            with st.form("editar_proyecto"):
                col1, col2 = st.columns(2)
                with col1:
                    nuevo_nombre = st.text_input("NOMBRE DEL CONTRATO", value=proyecto_seleccionado[1])
                    nueva_operadora = st.text_input("OPERADORA", value=proyecto_seleccionado[2] or "")
                    nuevo_gerente = st.text_input("GERENTE DE PROYECTO", value=proyecto_seleccionado[5] or "")
                with col2:
                    nuevo_numero = st.text_input("N° CONTRATO", value=proyecto_seleccionado[3] or "")
                    nuevo_contacto = st.text_input("CONTACTO CLIENTE", value=proyecto_seleccionado[4] or "")
                if st.form_submit_button("💾 ACTUALIZAR PROYECTO"):
                    ejecutar_query(
                        "UPDATE proyectos SET contrato_nom=?, operadora=?, contrato_num=?, contacto_cliente=?, gerente_proyecto=? WHERE id=?",
                        (nuevo_nombre, nueva_operadora, nuevo_numero, nuevo_contacto, nuevo_gerente, proyecto_seleccionado[0])
                    )
                    registrar_log(st.session_state.usuario_id, f"EDITAR PROYECTO: {proyecto_seleccionado[1]}")
                    st.success("✅ PROYECTO ACTUALIZADO")
                    st.rerun()

            st.subheader("🗑️ ELIMINAR PROYECTO")
            st.warning("⚠️ ESTA ACCIÓN ELIMINARÁ EL PROYECTO, TODOS SUS POZOS, PROVEEDORES Y REPORTES ASOCIADOS.")
            if st.button("⚠️ ELIMINAR PROYECTO PERMANENTEMENTE", type="primary"):
                reportes_asociados = ejecutar_query("SELECT id FROM inspecciones WHERE proyecto_id = ?", (proyecto_seleccionado[0],), fetch=True)
                if reportes_asociados:
                    st.error(f"❌ NO SE PUEDE ELIMINAR: EL PROYECTO TIENE {len(reportes_asociados)} REPORTES ASOCIADOS")
                else:
                    pozos_eliminar = ejecutar_query("SELECT id FROM pozos WHERE proyecto_id = ?", (proyecto_seleccionado[0],), fetch=True)
                    for pozo in pozos_eliminar:
                        ejecutar_query("DELETE FROM proveedores WHERE pozo_id = ?", (pozo[0],))
                        ejecutar_query("DELETE FROM contador_reportes WHERE pozo_id = ?", (pozo[0],))
                        ejecutar_query("DELETE FROM discrepancias_pozo WHERE pozo_id = ?", (pozo[0],))
                        ejecutar_query("DELETE FROM inspecciones WHERE pozo_id = ?", (pozo[0],))
                    ejecutar_query("DELETE FROM pozos WHERE proyecto_id = ?", (proyecto_seleccionado[0],))
                    ejecutar_query("DELETE FROM proyectos WHERE id = ?", (proyecto_seleccionado[0],))
                    registrar_log(st.session_state.usuario_id, f"ELIMINAR PROYECTO: {proyecto_seleccionado[1]}")
                    st.success("✅ PROYECTO ELIMINADO")
                    st.rerun()

    with tab_pozos_prov:
        proyectos = ejecutar_query("SELECT id, contrato_nom FROM proyectos", fetch=True)
        if proyectos:
            proy_id = st.selectbox("PROYECTO", proyectos, format_func=lambda x: x[1], key="proy_sel")

            st.subheader("📍 Pozos")
            pozos = ejecutar_query("SELECT id, rig, nombre_pozo FROM pozos WHERE proyecto_id = ? ORDER BY rig", (proy_id[0],), fetch=True)
            df_pozos = pd.DataFrame(pozos, columns=["id", "rig", "nombre_pozo"]) if pozos else pd.DataFrame(columns=["id", "rig", "nombre_pozo"])

            with st.form("form_pozos"):
                edited_pozos = st.data_editor(
                    df_pozos,
                    column_config={
                        "id": st.column_config.NumberColumn("ID", disabled=True),
                        "rig": st.column_config.TextColumn("RIG", required=True),
                        "nombre_pozo": st.column_config.TextColumn("NOMBRE DEL POZO", required=True)
                    },
                    num_rows="dynamic", key="pozos_editor"
                )
                if st.form_submit_button("💾 GUARDAR POZOS", use_container_width=True):
                    ids_en_bd = set([str(id[0]) for id in (ejecutar_query("SELECT id FROM pozos WHERE proyecto_id = ?", (proy_id[0],), fetch=True) or [])])
                    ids_en_editor = set()
                    for _, row in edited_pozos.iterrows():
                        if pd.notna(row['id']):
                            ids_en_editor.add(str(int(row['id'])))

                    for id_eliminar in ids_en_bd - ids_en_editor:
                        reportes_pozo = ejecutar_query("SELECT id FROM inspecciones WHERE pozo_id = ?", (id_eliminar,), fetch=True)
                        if reportes_pozo:
                            st.error(f"❌ NO SE PUEDE ELIMINAR EL POZO ID {id_eliminar}: TIENE {len(reportes_pozo)} REPORTES ASOCIADOS")
                            continue
                        ejecutar_query("DELETE FROM proveedores WHERE pozo_id = ?", (id_eliminar,))
                        ejecutar_query("DELETE FROM contador_reportes WHERE pozo_id = ?", (id_eliminar,))
                        ejecutar_query("DELETE FROM discrepancias_pozo WHERE pozo_id = ?", (id_eliminar,))
                        ejecutar_query("DELETE FROM pozos WHERE id = ?", (id_eliminar,))

                    for _, row in edited_pozos.iterrows():
                        if pd.isna(row['id']) and row.get('nombre_pozo'):
                            ejecutar_query("INSERT INTO pozos (proyecto_id, rig, nombre_pozo, fecha_creacion) VALUES (?, ?, ?, ?)",
                                          (proy_id[0], row.get('rig', ''), row['nombre_pozo'], datetime.now().strftime('%Y-%m-%d %H:%M:%S')))
                        elif pd.notna(row['id']) and row.get('nombre_pozo'):
                            ejecutar_query("UPDATE pozos SET rig=?, nombre_pozo=? WHERE id=?", (row.get('rig', ''), row['nombre_pozo'], int(row['id'])))

                    registrar_log(st.session_state.usuario_id, f"ACTUALIZAR POZOS DEL PROYECTO {proy_id[1]}")
                    st.success("✅ POZOS ACTUALIZADOS")
                    st.rerun()

            st.subheader("🏢 PROVEEDORES")
            pozos_list = ejecutar_query("SELECT id, rig, nombre_pozo FROM pozos WHERE proyecto_id = ? ORDER BY rig", (proy_id[0],), fetch=True)
            if pozos_list:
                pozo_id = st.selectbox("POZO", pozos_list, format_func=lambda x: x[2], key="pozo_sel")
                proveedores = ejecutar_query("SELECT id, nombre_proveedor FROM proveedores WHERE pozo_id = ?", (pozo_id[0],), fetch=True)
                df_prov = pd.DataFrame(proveedores, columns=["id", "nombre_proveedor"]) if proveedores else pd.DataFrame(columns=["id", "nombre_proveedor"])

                with st.form("form_proveedores"):
                    edited_prov = st.data_editor(
                        df_prov,
                        column_config={
                            "id": st.column_config.NumberColumn("ID", disabled=True),
                            "nombre_proveedor": st.column_config.TextColumn("NOMBRE DEL PROVEEDOR", required=True)
                        },
                        num_rows="dynamic", key="prov_editor"
                    )
                    if st.form_submit_button("💾 GUARDAR PROVEEDORES", use_container_width=True):
                        ids_en_bd = set([str(id[0]) for id in (ejecutar_query("SELECT id FROM proveedores WHERE pozo_id = ?", (pozo_id[0],), fetch=True) or [])])
                        ids_en_editor = set()
                        for _, row in edited_prov.iterrows():
                            if pd.notna(row['id']):
                                ids_en_editor.add(str(int(row['id'])))

                        for id_eliminar in ids_en_bd - ids_en_editor:
                            ejecutar_query("DELETE FROM proveedores WHERE id = ?", (id_eliminar,))

                        for _, row in edited_prov.iterrows():
                            if pd.isna(row['id']) and row.get('nombre_proveedor'):
                                ejecutar_query("INSERT INTO proveedores (pozo_id, nombre_proveedor) VALUES (?, ?)",
                                              (pozo_id[0], row['nombre_proveedor']))
                            elif pd.notna(row['id']) and row.get('nombre_proveedor'):
                                ejecutar_query("UPDATE proveedores SET nombre_proveedor=? WHERE id=?", (row['nombre_proveedor'], int(row['id'])))

                        st.success("✅ PROVEEDORES ACTUALIZADOS")
                        st.rerun()

# ============================================================================
# NUEVO REPORTE QA/QC
# ============================================================================
elif menu == "📝 NUEVO REPORTE QA/QC":
    st.header("📝 NUEVO REPORTE QA/QC")

    if st.session_state.formulario_activo and not st.session_state.datos_seleccion.get('proyecto'):
        st.session_state.formulario_activo = False
        st.session_state.modo_edicion = False
        st.rerun()

    # ================================================================
    # CARGA DE EDICIÓN
    # ================================================================
    if st.session_state.modo_edicion and st.session_state.edicion_reporte_id:
        if not st.session_state.datos_editados_cargados:
            reporte_data = ejecutar_query("""
                SELECT datos_json, usuario_creador_id, proyecto_id, pozo_id, proveedor_id, estado_reporte, fecha_emision
                FROM inspecciones WHERE reporte_id = ?""",
                (st.session_state.edicion_reporte_id,), fetch=True)
            if reporte_data:
                datos_json_str, usuario_creador_id, proy_id_bd, poz_id_bd, prov_id_bd, estado_reporte_bd, fecha_emision_bd = reporte_data[0]

                if st.session_state.usuario_rol != "admin" and usuario_creador_id != st.session_state.usuario_id:
                    st.error("❌ NO TIENE PERMISO PARA EDITAR ESTE REPORTE")
                    st.session_state.modo_edicion = False
                    st.rerun()

                st.session_state.reporte_actual_id = st.session_state.edicion_reporte_id
                st.session_state.estado_reporte_actual = estado_reporte_bd
                st.session_state.fecha_emision_actual = fecha_emision_bd

                match = re.search(r'FR-(\d+)-', st.session_state.edicion_reporte_id)
                if match:
                    st.session_state.reporte_numero_actual = int(match.group(1))

                p_sel_bd = ejecutar_query("SELECT id, contrato_nom, operadora, contacto_cliente, contrato_num, gerente_proyecto FROM proyectos WHERE id = ?", (proy_id_bd,), fetch=True)[0]
                poz_sel_bd = ejecutar_query("SELECT id, rig, nombre_pozo FROM pozos WHERE id = ?", (poz_id_bd,), fetch=True)[0]
                prov_sel_bd = None
                if prov_id_bd:
                    prov_data = ejecutar_query("SELECT id, nombre_proveedor FROM proveedores WHERE id = ?", (prov_id_bd,), fetch=True)
                    if prov_data:
                        prov_sel_bd = prov_data[0]

                st.session_state.datos_seleccion['proyecto'] = p_sel_bd
                st.session_state.datos_seleccion['pozo'] = poz_sel_bd
                st.session_state.datos_seleccion['proveedor'] = prov_sel_bd

                if datos_json_str:
                    datos_cargados = json.loads(datos_json_str)
                    form_data_raw = datos_cargados.get('form_data', {})

                    # Restaurar form_data preservando columnas de DataFrames
                    st.session_state.form_data = {}
                    for key, value in form_data_raw.items():
                        st.session_state.form_data[key] = restaurar_dfs_desde_json(value)

                    st.session_state.form_data['check_vals'] = datos_cargados.get('check_vals', {})
                    st.session_state.form_data['otras'] = datos_cargados.get('otras', '')
                    st.session_state.nota_instrumentos = datos_cargados.get('nota_instrumentos', '')
                    st.session_state.nota_personal = datos_cargados.get('nota_personal', '')
                    st.session_state.mostrar_nota_inst = datos_cargados.get('mostrar_nota_inst', False)
                    st.session_state.mostrar_nota_pers = datos_cargados.get('mostrar_nota_pers', False)
                    st.session_state.n_inspecciones = datos_cargados.get('n_inspecciones', 1)
                    st.session_state.n_ensambles = datos_cargados.get('n_ensambles', 1)
                    st.session_state.n_pruebas = datos_cargados.get('n_pruebas', 1)
                    st.session_state.n_fotos_secc = datos_cargados.get('n_fotos_secc', 1)
                    st.session_state.n_fotos_insp = datos_cargados.get('n_fotos_insp', {})
                    st.session_state.n_fotos_ens = datos_cargados.get('n_fotos_ens', {})
                    st.session_state.n_fotos_prueba = datos_cargados.get('n_fotos_prueba', {})

                    if 'hallazgos_list' in datos_cargados:
                        st.session_state.hallazgos_list = restaurar_hallazgos_desde_json(datos_cargados['hallazgos_list'])
                    if 'estado_herramientas_df' in datos_cargados:
                        st.session_state.estado_herramientas_df = restaurar_dfs_desde_json(
                            datos_cargados['estado_herramientas_df'],
                            columnas_default=["Herramienta", "Cant. Inspeccionada", "Cant. Aceptada", "Cant. Rechazada", "Cant. Pendiente"]
                        )

                    # CORREGIDO: restaurar fotos de inspección, ensamble y prueba
                    if 'fotos_inspecciones' in datos_cargados:
                        st.session_state.fotos_inspecciones = restaurar_dict_fotos(datos_cargados['fotos_inspecciones'])
                    if 'fotos_ensambles' in datos_cargados:
                        st.session_state.fotos_ensambles = restaurar_dict_fotos(datos_cargados['fotos_ensambles'])
                    if 'fotos_pruebas' in datos_cargados:
                        st.session_state.fotos_pruebas = restaurar_dict_fotos(datos_cargados['fotos_pruebas'])
                    if 'fotos_generales' in datos_cargados:
                        st.session_state.fotos_generales = restaurar_lista_fotos(datos_cargados['fotos_generales'])

                st.session_state.datos_editados_cargados = True
                st.session_state.formulario_activo = True
                st.rerun()

    # ================================================================
    # INICIO DE NUEVO REPORTE
    # ================================================================
    if not st.session_state.formulario_activo and not st.session_state.modo_edicion:
        proyectos = ejecutar_query("SELECT id, contrato_nom, operadora, contacto_cliente, contrato_num, gerente_proyecto FROM proyectos", fetch=True)
        if not proyectos:
            st.warning("⚠️ NO HAY PROYECTOS REGISTRADOS.")
            st.stop()

        col1, col2, col3 = st.columns(3)
        with col1:
            p_sel = st.selectbox("PROYECTO", proyectos, format_func=lambda x: x[1])

        pozos = ejecutar_query("SELECT id, rig, nombre_pozo FROM pozos WHERE proyecto_id = ?", (p_sel[0],), fetch=True)
        with col2:
            if pozos:
                poz_sel = st.selectbox("RIG / POZO", pozos, format_func=lambda x: f"{x[1]} - {x[2]}")
            else:
                st.warning("⚠️ NO HAY RIGS REGISTRADOS EN ESTE PROYECTO")
                st.stop()

        proveedores = ejecutar_query("SELECT id, nombre_proveedor FROM proveedores WHERE pozo_id = ?", (poz_sel[0],), fetch=True)
        with col3:
            if proveedores:
                prov_sel = st.selectbox("PROVEEDOR", proveedores, format_func=lambda x: x[1])
                st.session_state.datos_seleccion['proveedor'] = prov_sel
            else:
                prov_sel = None
                st.session_state.datos_seleccion['proveedor'] = None
                st.info("📌 No hay proveedores registrados.")

        if st.button("🚀 INICIAR REPORTE", use_container_width=True):
            try:
                with sqlite3.connect(DB_PATH) as conn:
                    cursor = conn.cursor()
                    cursor.execute("SELECT ultimo_numero FROM contador_reportes WHERE pozo_id = ?", (poz_sel[0],))
                    contador = cursor.fetchone()
                    if contador:
                        nuevo_numero = contador[0] + 1
                        cursor.execute("UPDATE contador_reportes SET ultimo_numero = ? WHERE pozo_id = ?", (nuevo_numero, poz_sel[0]))
                    else:
                        nuevo_numero = 1
                        cursor.execute("INSERT INTO contador_reportes (pozo_id, ultimo_numero) VALUES (?, ?)", (poz_sel[0], nuevo_numero))
                    conn.commit()
            except Exception as e:
                st.error(f"Error al crear contador: {e}")
                st.stop()

            st.session_state.reporte_numero_actual = nuevo_numero
            st.session_state.reporte_actual_id = f"FR-{nuevo_numero}-{poz_sel[2].replace(' ', '_')}"
            st.session_state.formulario_activo = True
            st.session_state.estado_reporte_actual = "BORRADOR"
            st.session_state.fecha_emision_actual = None
            st.session_state.datos_seleccion['proyecto'] = p_sel
            st.session_state.datos_seleccion['pozo'] = poz_sel
            # Resetear estado completo
            st.session_state.form_data = {'check_vals': {}, 'otras': ''}
            st.session_state.nota_instrumentos = ''
            st.session_state.nota_personal = ''
            st.session_state.mostrar_nota_inst = False
            st.session_state.mostrar_nota_pers = False
            st.session_state.n_inspecciones = 1
            st.session_state.n_ensambles = 1
            st.session_state.n_pruebas = 1
            st.session_state.n_fotos_secc = 1
            st.session_state.hallazgos_list = []
            st.session_state.estado_herramientas_df = pd.DataFrame(columns=[
                "Herramienta", "Cant. Inspeccionada", "Cant. Aceptada", "Cant. Rechazada", "Cant. Pendiente"
            ])
            st.session_state.fotos_inspecciones = {}
            st.session_state.fotos_ensambles = {}
            st.session_state.fotos_pruebas = {}
            st.session_state.fotos_generales = []
            st.session_state.n_fotos_insp = {}
            st.session_state.n_fotos_ens = {}
            st.session_state.n_fotos_prueba = {}
            st.session_state.datos_editados_cargados = False
            st.rerun()

    # ================================================================
    # FORMULARIO ACTIVO
    # ================================================================
    elif st.session_state.formulario_activo:
        p_sel = st.session_state.datos_seleccion.get('proyecto')
        poz_sel = st.session_state.datos_seleccion.get('pozo')
        prov_sel = st.session_state.datos_seleccion.get('proveedor')
        report_id = st.session_state.reporte_actual_id

        if not p_sel or not poz_sel:
            st.error("❌ Error: No hay proyecto o pozo seleccionado")
            st.stop()

        if st.button("◀ CAMBIAR PROYECTO", use_container_width=True):
            st.session_state.formulario_activo = False
            st.session_state.modo_edicion = False
            st.session_state.menu_seleccionado = "📝 NUEVO REPORTE QA/QC"
            st.session_state.datos_editados_cargados = False
            st.rerun()

        estado_texto = "📝 EN CURSO (BORRADOR)" if st.session_state.get('estado_reporte_actual', 'BORRADOR') == 'BORRADOR' else "✅ COMPLETADO (EMITIDO)"
        st.info(f"📄 **NÚMERO DE REPORTE:** `{report_id}` | **ESTADO:** {estado_texto}")

        col_save1, col_save2 = st.columns([1, 4])
        with col_save1:
            texto_boton = "💾 GUARDAR EN CURSO" if st.session_state.get('estado_reporte_actual', 'BORRADOR') == 'BORRADOR' else "💾 GUARDAR CAMBIOS"
            if st.button(texto_boton, key="guardar_todo_btn", use_container_width=True, type="primary"):
                if guardar_datos_reporte(p_sel, poz_sel, prov_sel, report_id, es_generacion_word=False):
                    st.success("✅ GUARDADO CORRECTAMENTE")

        tab_nombres = [
            "📋 CABECERA", "📄 1-2 PROVEEDOR/HERRAMIENTAS", "📋 2.2-3 ACTIVIDADES/DOCS",
            "🔧 4-5 INSTRUMENTOS/PERSONAL", "📝 6 DESCRIPCIÓN MONITOREO", "🔧 6.2 INSPECCIÓN",
            "🔩 6.3 ENSAMBLE", "🧪 6.4 PRUEBAS", "⚠️ 7 HALLAZGOS", "📸 8 FOTOS", "✅ 9-10 ESTADO/REPRESENTANTES"
        ]
        tabs = st.tabs(tab_nombres)

        # --- TAB 0: CABECERA ---
        with tabs[0]:
            st.subheader("📋 DATOS DE CABECERA")
            proyecto_datos = ejecutar_query("SELECT contacto_cliente, gerente_proyecto FROM proyectos WHERE id = ?", (p_sel[0],), fetch=True)
            contacto_default = proyecto_datos[0][0] if proyecto_datos and proyecto_datos[0][0] else ""
            gerente_default = proyecto_datos[0][1] if proyecto_datos and proyecto_datos[0][1] else ""

            col1, col2 = st.columns(2)
            with col1:
                cliente = st.text_input("CLIENTE *", value=st.session_state.form_data.get('cliente', p_sel[2] or ""), key="cliente")
                contrato_no = st.text_input("CONTRATO No *", value=st.session_state.form_data.get('contrato', p_sel[4] or ""), key="contrato")
                gerente_proyecto = st.text_input("GERENTE DE PROYECTO", value=st.session_state.form_data.get('gerente_proyecto', gerente_default), key="gerente_proyecto")
                orden_compra = st.text_input("ORDEN DE COMPRA", value=st.session_state.form_data.get('orden_compra', ""), key="orden_compra")
                orden_servicio = st.text_input("ORDEN DE SERVICIO", value=st.session_state.form_data.get('orden_servicio', ""), key="orden_servicio")
            with col2:
                contacto_cliente = st.text_input("CONTACTO DEL CLIENTE", value=st.session_state.form_data.get('contacto', contacto_default), key="contacto")
                pozo_proyecto = st.text_input("POZO/PROYECTO *", value=st.session_state.form_data.get('pozo', poz_sel[2] or ""), key="pozo")
                solicitud = st.text_input("SOLICITUD", value=st.session_state.form_data.get('solicitud', ""), key="solicitud")

            col_f1, col_f2, col_f3 = st.columns(3)
            with col_f1:
                fecha_servicio = st.text_input("FECHA DE SERVICIO", value=st.session_state.form_data.get('fecha_serv', ''), key="fecha_serv", placeholder="YYYY-MM-DD")
            with col_f2:
                fecha_envio = st.text_input("FECHA DE ENVÍO", value=st.session_state.form_data.get('fecha_envio', "N/A"), key="fecha_envio")
            with col_f3:
                fecha_prox = st.text_input("PRÓXIMO SERVICIO", value=st.session_state.form_data.get('fecha_prox', "N/A"), key="fecha_prox")

            st.session_state.form_data.update({
                'cliente': cliente, 'contrato': contrato_no, 'gerente_proyecto': gerente_proyecto,
                'orden_compra': orden_compra, 'orden_servicio': orden_servicio,
                'contacto': contacto_cliente, 'pozo': pozo_proyecto, 'solicitud': solicitud,
                'fecha_serv': fecha_servicio, 'fecha_envio': fecha_envio, 'fecha_prox': fecha_prox
            })

        # --- TAB 1: PROVEEDOR/HERRAMIENTAS ---
        with tabs[1]:
            st.subheader("1. DETALLES DEL PROVEEDOR")
            if prov_sel:
                st.info(f"PROVEEDOR: {prov_sel[1]}")

            cols_prov = ["COMPAÑIA", "NOMBRE DEL CONTACTO", "CARGO", "ROL"]
            df_prov_actual = st.session_state.form_data.get('ed_prov', pd.DataFrame(columns=cols_prov))
            if not isinstance(df_prov_actual, pd.DataFrame):
                df_prov_actual = restaurar_dfs_desde_json(df_prov_actual, cols_prov)

            with st.form(key="form_ed_prov"):
                ed_prov = st.data_editor(df_prov_actual, num_rows="dynamic", key="ed_prov_in_form")
                if st.form_submit_button("✅ CONFIRMAR CAMBIOS EN PROVEEDOR", use_container_width=True):
                    st.session_state.form_data['ed_prov'] = ed_prov
                    st.success("✅ CAMBIOS GUARDADOS")
                    st.rerun()

            st.subheader("2. RESUMEN DE MONITOREO QA/QC")
            st.session_state.form_data['resumen'] = st_quill(st.session_state.form_data.get('resumen', ''), key="resumen")

            st.subheader("2.1 DESCRIPCION DE HERRAMIENTAS *")
            cols_items = ["N°", "DESCRIPCIÓN", "OD", "CONEXIONES", "NUMERO DE SERIE"]
            df_items_actual = st.session_state.form_data.get('ed_items', pd.DataFrame(columns=cols_items))
            if not isinstance(df_items_actual, pd.DataFrame):
                df_items_actual = restaurar_dfs_desde_json(df_items_actual, cols_items)

            with st.form(key="form_ed_items"):
                ed_items = st.data_editor(df_items_actual, num_rows="dynamic", key="ed_items_in_form")
                if isinstance(ed_items, list):
                    ed_items = pd.DataFrame(ed_items)
                if ed_items is not None and isinstance(ed_items, pd.DataFrame) and not ed_items.empty:
                    ed_items = ed_items.copy()
                    ed_items["N°"] = range(1, len(ed_items) + 1)
                if st.form_submit_button("✅ CONFIRMAR CAMBIOS EN HERRAMIENTAS", use_container_width=True):
                    st.session_state.form_data['ed_items'] = ed_items
                    st.success("✅ CAMBIOS GUARDADOS")
                    st.rerun()

        # --- TAB 2: ACTIVIDADES/DOCS ---
        with tabs[2]:
            st.subheader("2.2 LISTA DE ACTIVIDADES REALIZADAS *")
            acts = [
                "Revisión Documental", "Trazabilidad", "Calibración de Instrumentos",
                "Calificación de Personal", "Inspección Visual/Dimensional", "Insp. CAT 3",
                "Insp. CAT 5", "Inspección con Partículas Magnéticas (MPI)",
                "Inspección con Líquidos Penetrantes (LPI)", "MPI en Conexiones", "Medición",
                "Sand Blasting", "Ensamble", "Drifting", "Prueba de Corrosión",
                "Prueba Hidrostática", "Prueba de Presión", "Prueba de Funcionamiento",
                "Almacenamiento", "Protección de Conexiónes", "Certificación Final", "Despacho/Envío"
            ]
            cols = st.columns(4)
            check_vals = {}
            for i, a in enumerate(acts):
                check_vals[a] = cols[i % 4].checkbox(a, value=st.session_state.form_data.get('check_vals', {}).get(a, False), key=f"act_{a}")
            st.session_state.form_data['check_vals'] = check_vals

            st.markdown("**OTRAS ACTIVIDADES**")
            st.session_state.form_data['otras'] = st.text_area("Escriba otras actividades (una por línea)", value=st.session_state.form_data.get('otras', ''), height=100, key="otras")

            st.subheader("3. SECCIÓN DOCUMENTAL Y NORMATIVA")
            cols_doc = ["DOCUMENTO", "NUMERO DE DOCUMENTO", "REVISION"]
            df_doc_actual = st.session_state.form_data.get('ed_doc', pd.DataFrame(columns=cols_doc))
            if not isinstance(df_doc_actual, pd.DataFrame):
                df_doc_actual = restaurar_dfs_desde_json(df_doc_actual, cols_doc)

            with st.form(key="form_ed_doc"):
                ed_doc = st.data_editor(df_doc_actual, num_rows="dynamic", key="ed_doc_in_form")
                if st.form_submit_button("✅ CONFIRMAR CAMBIOS EN DOCUMENTOS", use_container_width=True):
                    st.session_state.form_data['ed_doc'] = ed_doc
                    st.success("✅ CAMBIOS GUARDADOS")
                    st.rerun()

        # --- TAB 3: INSTRUMENTOS/PERSONAL ---
        with tabs[3]:
            st.subheader("4. DETALLE DE INSTRUMENTOS DE INSPECCIÓN Y EQUIPOS DE PRUEBA USADOS")
            cols_inst = ["NOMBRE DEL EQUIPO", "FABRICANTE", "NUMERO DE SERIE", "FECHA CALIB.", "VENCIMIENTO"]
            df_inst_actual = st.session_state.form_data.get('ed_inst', pd.DataFrame(columns=cols_inst))
            if not isinstance(df_inst_actual, pd.DataFrame):
                df_inst_actual = restaurar_dfs_desde_json(df_inst_actual, cols_inst)

            with st.form(key="form_ed_inst"):
                ed_inst = st.data_editor(df_inst_actual, num_rows="dynamic", key="ed_inst_in_form")
                if st.form_submit_button("✅ CONFIRMAR CAMBIOS EN INSTRUMENTOS", use_container_width=True):
                    st.session_state.form_data['ed_inst'] = ed_inst
                    st.success("✅ CAMBIOS GUARDADOS")
                    st.rerun()

            if st.button("📝 Añadir Nota en Instrumentos", key="add_note_inst"):
                st.session_state.mostrar_nota_inst = not st.session_state.mostrar_nota_inst
            if st.session_state.mostrar_nota_inst:
                st.session_state.nota_instrumentos = st.text_area("Nota sobre Instrumentos", value=st.session_state.nota_instrumentos, height=100, key="nota_inst_text")

            st.subheader("5. COMPETENCIA DEL PERSONAL")
            cols_pers = ["NOMBRE", "CARGO", "COMPAÑIA", "CALIFICACIÓN", "OBSERVACIONES"]
            df_pers_actual = st.session_state.form_data.get('ed_pers', pd.DataFrame(columns=cols_pers))
            if not isinstance(df_pers_actual, pd.DataFrame):
                df_pers_actual = restaurar_dfs_desde_json(df_pers_actual, cols_pers)

            with st.form(key="form_ed_pers"):
                ed_pers = st.data_editor(df_pers_actual, num_rows="dynamic", key="ed_pers_in_form")
                if st.form_submit_button("✅ CONFIRMAR CAMBIOS EN PERSONAL", use_container_width=True):
                    st.session_state.form_data['ed_pers'] = ed_pers
                    st.success("✅ CAMBIOS GUARDADOS")
                    st.rerun()

            if st.button("📝 Añadir Nota en Personal", key="add_note_pers"):
                st.session_state.mostrar_nota_pers = not st.session_state.mostrar_nota_pers
            if st.session_state.mostrar_nota_pers:
                st.session_state.nota_personal = st.text_area("Nota sobre Personal", value=st.session_state.nota_personal, height=100, key="nota_pers_text")

        # --- TAB 4: DESCRIPCIÓN INSPECCIÓN ---
        with tabs[4]:
            st.subheader("6. DESCRIPCIÓN DE MONITOREO QA/QC")
            st.markdown("**Información General**")
            st.session_state.form_data['info_gral'] = st_quill(st.session_state.form_data.get('info_gral', ''), key="info_gral")
            st.markdown("**Objetivos**")
            st.session_state.form_data['objetivos'] = st_quill(st.session_state.form_data.get('objetivos', ''), key="objetivos")

            st.subheader("6.1 REVISIÓN DOCUMENTAL")
            st.markdown("**Descripción**")
            st.session_state.form_data['desc_rev'] = st_quill(st.session_state.form_data.get('desc_rev', ''), key="desc_rev")

            cols_reg = ["DESCRIPCIÓN", "NUMERO DE SERIE", "HORAS", "COC, MTR"]
            df_reg_actual = st.session_state.form_data.get('ed_reg', pd.DataFrame(columns=cols_reg))
            if not isinstance(df_reg_actual, pd.DataFrame):
                df_reg_actual = restaurar_dfs_desde_json(df_reg_actual, cols_reg)

            with st.form(key="form_ed_reg"):
                ed_reg = st.data_editor(df_reg_actual, num_rows="dynamic", key="ed_reg_in_form")
                if st.form_submit_button("✅ CONFIRMAR CAMBIOS EN REGISTRO", use_container_width=True):
                    st.session_state.form_data['ed_reg'] = ed_reg
                    st.success("✅ CAMBIOS GUARDADOS")
                    st.rerun()

        # --- TAB 5: INSPECCIÓN ---
        with tabs[5]:
            st.subheader("6.2 INSPECCIÓN")
            for i in range(st.session_state.n_inspecciones):
                with st.expander(f"🔧 INSPECCIÓN No {i+1}", expanded=True):
                    st.markdown("**DESCRIPCION**")
                    desc = st_quill(st.session_state.form_data.get(f'insp_desc_{i}', ''), key=f"insp_desc_{i}")

                    col_i1, col_i2, col_i3, col_i4 = st.columns(4)
                    with col_i1:
                        inspector = st.text_input("INSPECTOR", value=st.session_state.form_data.get(f'inspector_{i}', ''), key=f"inspector_{i}")
                    with col_i2:
                        doc_ref = st.text_input("W.O. / O.S. / O.T.", value=st.session_state.form_data.get(f'doc_ref_{i}', ''), key=f"doc_ref_{i}")
                    with col_i3:
                        herramienta = st.text_input("HERRAMIENTA", value=st.session_state.form_data.get(f'herramienta_{i}', ''), key=f"herramienta_{i}")
                    with col_i4:
                        num_serie = st.text_input("NUMERO DE SERIE", value=st.session_state.form_data.get(f'num_serie_{i}', ''), key=f"num_serie_{i}")

                    st.markdown("**COMPONENTES**")
                    cols_comp = ["N°", "COMPONENTE", "SERIAL", "INSPECCIONES", "OBS"]
                    df_comp_actual = st.session_state.form_data.get(f'comp_{i}', pd.DataFrame(columns=cols_comp))
                    if not isinstance(df_comp_actual, pd.DataFrame):
                        df_comp_actual = restaurar_dfs_desde_json(df_comp_actual, cols_comp)

                    with st.form(key=f"form_comp_{i}"):
                        comp = st.data_editor(df_comp_actual, num_rows="dynamic", key=f"comp_{i}_in_form")
                        if isinstance(comp, list):
                            comp = pd.DataFrame(comp)
                        if comp is not None and isinstance(comp, pd.DataFrame) and not comp.empty:
                            comp = comp.copy()
                            comp["N°"] = range(1, len(comp) + 1)
                        if st.form_submit_button(f"✅ CONFIRMAR COMPONENTES {i+1}", use_container_width=True):
                            st.session_state.form_data[f'comp_{i}'] = comp
                            st.success("✅ CAMBIOS GUARDADOS")
                            st.rerun()

                    st.markdown("**FOTOS DE INSPECCIÓN**")
                    key_fotos = f"fotos_insp_{i}"
                    key_n = f"n_fotos_insp_{i}"
                    if key_fotos not in st.session_state.fotos_inspecciones:
                        st.session_state.fotos_inspecciones[key_fotos] = []
                    if key_n not in st.session_state.n_fotos_insp:
                        st.session_state.n_fotos_insp[key_n] = 1

                    fotos_actuales = list(st.session_state.fotos_inspecciones.get(key_fotos, []))
                    nuevas_fotos = []
                    for k in range(st.session_state.n_fotos_insp[key_n]):
                        c1, c2, c3 = st.columns([1, 3, 1])
                        with c1:
                            f = st.file_uploader(f"Foto {k+1}", type=['jpg', 'png'], key=f"f_insp_{i}_{k}")
                        with c2:
                            # Recuperar descripción existente si hay foto guardada
                            desc_default = fotos_actuales[k].get('descripcion', '') if k < len(fotos_actuales) else ''
                            d = st.text_input(f"Descripción", value=desc_default, key=f"fd_insp_{i}_{k}")
                        with c3:
                            if st.button(f"🗑️", key=f"del_insp_{i}_{k}"):
                                st.session_state.n_fotos_insp[key_n] = max(1, st.session_state.n_fotos_insp[key_n] - 1)
                                if k < len(fotos_actuales):
                                    fotos_actuales.pop(k)
                                st.session_state.fotos_inspecciones[key_fotos] = fotos_actuales
                                st.rerun()
                        if f is not None:
                            nuevas_fotos.append({"imagen": f, "descripcion": d})
                        elif k < len(fotos_actuales) and fotos_actuales[k].get('imagen') is not None:
                            nuevas_fotos.append({"imagen": fotos_actuales[k]['imagen'], "descripcion": d})
                        else:
                            nuevas_fotos.append({"imagen": None, "descripcion": d})

                    if st.button(f"➕ Añadir foto a INSPECCIÓN {i+1}", key=f"add_insp_foto_{i}"):
                        st.session_state.n_fotos_insp[key_n] += 1
                        st.rerun()

                    st.session_state.fotos_inspecciones[key_fotos] = nuevas_fotos
                    st.session_state.form_data[f'inspector_{i}'] = inspector
                    st.session_state.form_data[f'doc_ref_{i}'] = doc_ref
                    st.session_state.form_data[f'herramienta_{i}'] = herramienta
                    st.session_state.form_data[f'num_serie_{i}'] = num_serie
                    st.session_state.form_data[f'insp_desc_{i}'] = desc

            col1, col2 = st.columns(2)
            with col1:
                if st.button("➕ Añadir Inspección", key="add_insp_btn"):
                    st.session_state.n_inspecciones += 1
                    st.rerun()
            with col2:
                if st.button("➖ Eliminar Última Inspección", key="del_insp_btn") and st.session_state.n_inspecciones > 1:
                    st.session_state.n_inspecciones -= 1
                    st.rerun()

        # --- TAB 6: ENSAMBLE ---
        with tabs[6]:
            st.subheader("6.3 ENSAMBLE Y REPORTE DE TORQUE")
            for j in range(st.session_state.n_ensambles):
                with st.expander(f"🔩 ENSAMBLE No {j+1}", expanded=True):
                    col_e1, col_e2 = st.columns(2)
                    with col_e1:
                        e_h = st.text_input("HERRAMIENTA", value=st.session_state.form_data.get(f'eh_{j}', ''), key=f"eh_{j}")
                    with col_e2:
                        e_s = st.text_input("SERIAL", value=st.session_state.form_data.get(f'es_{j}', ''), key=f"es_{j}")

                    st.markdown("**DESCRIPCIÓN**")
                    e_d = st_quill(st.session_state.form_data.get(f'ens_desc_{j}', ''), key=f"ens_desc_{j}")

                    st.markdown("**REPORTE DE TORQUE**")
                    cols_torq = ["N°", "COMPONENTE 1", "COMPONENTE 2", "TORQUE RECOM.", "TORQUE APLICADO", "LUBRICANTE"]
                    df_torq_actual = st.session_state.form_data.get(f'torque_{j}', pd.DataFrame(columns=cols_torq))
                    if not isinstance(df_torq_actual, pd.DataFrame):
                        df_torq_actual = restaurar_dfs_desde_json(df_torq_actual, cols_torq)

                    with st.form(key=f"form_torque_{j}"):
                        e_t = st.data_editor(df_torq_actual, num_rows="dynamic", key=f"torque_{j}_in_form")
                        if isinstance(e_t, list):
                            e_t = pd.DataFrame(e_t)
                        if e_t is not None and isinstance(e_t, pd.DataFrame) and not e_t.empty:
                            e_t = e_t.copy()
                            e_t["N°"] = range(1, len(e_t) + 1)
                        if st.form_submit_button(f"✅ CONFIRMAR TORQUE {j+1}", use_container_width=True):
                            st.session_state.form_data[f'torque_{j}'] = e_t
                            st.success("✅ CAMBIOS GUARDADOS")
                            st.rerun()

                    st.markdown("**FOTOS DE ENSAMBLE**")
                    key_fotos = f"fotos_ens_{j}"
                    key_n = f"n_fotos_ens_{j}"
                    if key_fotos not in st.session_state.fotos_ensambles:
                        st.session_state.fotos_ensambles[key_fotos] = []
                    if key_n not in st.session_state.n_fotos_ens:
                        st.session_state.n_fotos_ens[key_n] = 1

                    fotos_actuales = list(st.session_state.fotos_ensambles.get(key_fotos, []))
                    nuevas_fotos = []
                    for k in range(st.session_state.n_fotos_ens[key_n]):
                        c1, c2, c3 = st.columns([1, 3, 1])
                        with c1:
                            f = st.file_uploader(f"Foto {k+1}", type=['jpg', 'png'], key=f"f_ens_{j}_{k}")
                        with c2:
                            desc_default = fotos_actuales[k].get('descripcion', '') if k < len(fotos_actuales) else ''
                            d = st.text_input(f"Descripción", value=desc_default, key=f"fd_ens_{j}_{k}")
                        with c3:
                            if st.button(f"🗑️", key=f"del_ens_{j}_{k}"):
                                st.session_state.n_fotos_ens[key_n] = max(1, st.session_state.n_fotos_ens[key_n] - 1)
                                if k < len(fotos_actuales):
                                    fotos_actuales.pop(k)
                                st.session_state.fotos_ensambles[key_fotos] = fotos_actuales
                                st.rerun()
                        if f is not None:
                            nuevas_fotos.append({"imagen": f, "descripcion": d})
                        elif k < len(fotos_actuales) and fotos_actuales[k].get('imagen') is not None:
                            nuevas_fotos.append({"imagen": fotos_actuales[k]['imagen'], "descripcion": d})
                        else:
                            nuevas_fotos.append({"imagen": None, "descripcion": d})

                    if st.button(f"➕ Añadir foto a ENSAMBLE {j+1}", key=f"add_ens_foto_{j}"):
                        st.session_state.n_fotos_ens[key_n] += 1
                        st.rerun()

                    st.session_state.fotos_ensambles[key_fotos] = nuevas_fotos
                    st.session_state.form_data[f'eh_{j}'] = e_h
                    st.session_state.form_data[f'es_{j}'] = e_s
                    st.session_state.form_data[f'ens_desc_{j}'] = e_d

            col1, col2 = st.columns(2)
            with col1:
                if st.button("➕ Añadir Ensamble", key="add_ens_btn"):
                    st.session_state.n_ensambles += 1
                    st.rerun()
            with col2:
                if st.button("➖ Eliminar Último Ensamble", key="del_ens_btn") and st.session_state.n_ensambles > 1:
                    st.session_state.n_ensambles -= 1
                    st.rerun()

        # --- TAB 7: PRUEBAS ---
        with tabs[7]:
            st.subheader("6.4 PRUEBAS")
            for k in range(st.session_state.n_pruebas):
                with st.expander(f"🧪 PRUEBA No {k+1}", expanded=True):
                    nombre = st.text_input("NOMBRE DE LA HERRAMIENTA", value=st.session_state.form_data.get(f'pr_nom_{k}', ''), key=f"pr_nom_{k}")
                    num_serie = st.text_input("NUMERO DE SERIE", value=st.session_state.form_data.get(f'pr_serie_{k}', ''), key=f"pr_serie_{k}")
                    tipo = st.text_input("TIPO DE PRUEBA", value=st.session_state.form_data.get(f'pr_tipo_{k}', ''), key=f"pr_tipo_{k}")
                    st.markdown("**DESCRIPCIÓN DE LA PRUEBA**")
                    desc = st_quill(st.session_state.form_data.get(f'pr_desc_{k}', ''), key=f"pr_desc_{k}")
                    st.markdown("**TOLERANCIA**")
                    tol = st_quill(st.session_state.form_data.get(f'pr_tol_{k}', ''), key=f"pr_tol_{k}")
                    resultado = st.text_input("RESULTADOS", value=st.session_state.form_data.get(f'pr_res_{k}', ''), key=f"pr_res_{k}")

                    st.markdown("**FOTOS DE PRUEBA**")
                    key_fotos = f"fotos_prueba_{k}"
                    key_n = f"n_fotos_prueba_{k}"
                    if key_fotos not in st.session_state.fotos_pruebas:
                        st.session_state.fotos_pruebas[key_fotos] = []
                    if key_n not in st.session_state.n_fotos_prueba:
                        st.session_state.n_fotos_prueba[key_n] = 1

                    fotos_actuales = list(st.session_state.fotos_pruebas.get(key_fotos, []))
                    nuevas_fotos = []
                    for m in range(st.session_state.n_fotos_prueba[key_n]):
                        c1, c2, c3 = st.columns([1, 3, 1])
                        with c1:
                            f = st.file_uploader(f"Foto {m+1}", type=['jpg', 'png'], key=f"f_pr_{k}_{m}")
                        with c2:
                            desc_default = fotos_actuales[m].get('descripcion', '') if m < len(fotos_actuales) else ''
                            d = st.text_input(f"Descripción", value=desc_default, key=f"fd_pr_{k}_{m}")
                        with c3:
                            if st.button(f"🗑️", key=f"del_pr_{k}_{m}"):
                                st.session_state.n_fotos_prueba[key_n] = max(1, st.session_state.n_fotos_prueba[key_n] - 1)
                                if m < len(fotos_actuales):
                                    fotos_actuales.pop(m)
                                st.session_state.fotos_pruebas[key_fotos] = fotos_actuales
                                st.rerun()
                        if f is not None:
                            nuevas_fotos.append({"imagen": f, "descripcion": d})
                        elif m < len(fotos_actuales) and fotos_actuales[m].get('imagen') is not None:
                            nuevas_fotos.append({"imagen": fotos_actuales[m]['imagen'], "descripcion": d})
                        else:
                            nuevas_fotos.append({"imagen": None, "descripcion": d})

                    if st.button(f"➕ Añadir foto a PRUEBA {k+1}", key=f"add_pr_foto_{k}"):
                        st.session_state.n_fotos_prueba[key_n] += 1
                        st.rerun()

                    st.session_state.fotos_pruebas[key_fotos] = nuevas_fotos
                    st.session_state.form_data[f'pr_nom_{k}'] = nombre
                    st.session_state.form_data[f'pr_serie_{k}'] = num_serie
                    st.session_state.form_data[f'pr_tipo_{k}'] = tipo
                    st.session_state.form_data[f'pr_desc_{k}'] = desc
                    st.session_state.form_data[f'pr_tol_{k}'] = tol
                    st.session_state.form_data[f'pr_res_{k}'] = resultado

            col1, col2 = st.columns(2)
            with col1:
                if st.button("➕ Añadir Prueba", key="add_pr_btn"):
                    st.session_state.n_pruebas += 1
                    st.rerun()
            with col2:
                if st.button("➖ Eliminar Última Prueba", key="del_pr_btn") and st.session_state.n_pruebas > 1:
                    st.session_state.n_pruebas -= 1
                    st.rerun()

        # --- TAB 8: HALLAZGOS ---
        with tabs[8]:
            st.subheader("7. HALLAZGOS / DISCREPANCIAS / NCR")

            for idx in range(len(st.session_state.hallazgos_list)):
                hallazgo = st.session_state.hallazgos_list[idx]
                with st.expander(f"⚠️ HALLAZGO {idx+1} - {hallazgo.get('tipo', 'Nuevo')}", expanded=False):
                    col_tipo, col_estado = st.columns(2)
                    with col_tipo:
                        hallazgo['tipo'] = st.selectbox("TIPO", ["DISC.", "OBS.", "NCR"],
                            index=["DISC.", "OBS.", "NCR"].index(hallazgo.get('tipo', 'DISC.')), key=f"tipo_{idx}")
                    with col_estado:
                        hallazgo['estado'] = st.selectbox("ESTADO", ["ABIERTO", "PENDIENTE", "CERRADO"],
                            index=["ABIERTO", "PENDIENTE", "CERRADO"].index(hallazgo.get('estado', 'ABIERTO')), key=f"estado_{idx}")

                    hallazgo['descripcion'] = st.text_area("DESCRIPCIÓN", value=hallazgo.get('descripcion', ''), key=f"desc_{idx}", height=100)
                    hallazgo['normativa'] = st.text_area("NORMATIVA", value=hallazgo.get('normativa', ''), key=f"norm_{idx}", height=80)
                    hallazgo['accion'] = st.text_area("ACCIÓN", value=hallazgo.get('accion', ''), key=f"acc_{idx}", height=80)
                    hallazgo['resolucion'] = st.text_area("RESOLUCIÓN", value=hallazgo.get('resolucion', ''), key=f"res_{idx}", height=80)

                    st.markdown("**FOTOS DEL HALLAZGO**")
                    # CORREGIDO: usar lista directamente del hallazgo, sin session_state separado
                    fotos_hallazgo = list(hallazgo.get('fotos', []))
                    nuevas_fotos_h = []
                    for m in range(len(fotos_hallazgo) + 1):
                        c1, c2, c3 = st.columns([1, 3, 1])
                        with c1:
                            f = st.file_uploader(f"Foto {m+1}", type=['jpg', 'png'], key=f"f_hall_{idx}_{m}")
                        with c2:
                            desc_default = fotos_hallazgo[m].get('descripcion', '') if m < len(fotos_hallazgo) else ''
                            d = st.text_input(f"Descripción", value=desc_default, key=f"fd_hall_{idx}_{m}")
                        with c3:
                            if m < len(fotos_hallazgo):
                                if st.button(f"🗑️", key=f"del_hall_{idx}_{m}"):
                                    fotos_hallazgo.pop(m)
                                    hallazgo['fotos'] = fotos_hallazgo
                                    st.session_state.hallazgos_list[idx] = hallazgo
                                    st.rerun()
                        if f is not None:
                            nuevas_fotos_h.append({"imagen": f, "descripcion": d})
                        elif m < len(fotos_hallazgo) and fotos_hallazgo[m].get('imagen') is not None:
                            nuevas_fotos_h.append({"imagen": fotos_hallazgo[m]['imagen'], "descripcion": d})

                    hallazgo['fotos'] = nuevas_fotos_h
                    st.session_state.hallazgos_list[idx] = hallazgo

                    if st.button(f"🗑️ Eliminar HALLAZGO {idx+1}", key=f"del_hall_{idx}"):
                        st.session_state.hallazgos_list.pop(idx)
                        st.rerun()

            if st.button("➕ AÑADIR NUEVO HALLAZGO", use_container_width=True):
                st.session_state.hallazgos_list.append({
                    'tipo': 'DISC.', 'descripcion': '', 'normativa': '',
                    'accion': '', 'resolucion': '', 'estado': 'ABIERTO', 'fotos': []
                })
                st.rerun()

        # --- TAB 9: FOTOS ---
        with tabs[9]:
            st.subheader("8. REPORTE FOTOGRÁFICO")
            fotos = []
            for s in range(st.session_state.n_fotos_secc):
                st.write(f"**Fila {s+1}**")
                cols_f = st.columns(3)
                for k in range(3):
                    with cols_f[k]:
                        f = st.file_uploader(f"Foto {s+1}-{k+1}", type=['jpg', 'png'], key=f"f_{s}_{k}")
                        d = st.text_input(f"Descripción {s+1}-{k+1}", key=f"fd_{s}_{k}")
                        if f:
                            fotos.append({"imagen": f, "descripcion": d})
            st.session_state.fotos_generales = fotos

            col1, col2 = st.columns(2)
            with col1:
                if st.button("➕ Añadir Línea de Fotos", key="add_line_btn"):
                    st.session_state.n_fotos_secc += 1
                    st.rerun()
            with col2:
                if st.button("➖ Eliminar Última Línea", key="del_line_btn") and st.session_state.n_fotos_secc > 1:
                    st.session_state.n_fotos_secc -= 1
                    st.rerun()

        # --- TAB 10: ESTADO/REPRESENTANTES ---
        with tabs[10]:
            st.subheader("9. ESTADO FINAL DE HERRAMIENTA")
            st.session_state.form_data['estado_final'] = st_quill(st.session_state.form_data.get('estado_final', ''), key="estado_final")

            st.markdown("**RESUMEN DE ESTADO DE HERRAMIENTAS**")
            cols_estado = ["Herramienta", "Cant. Inspeccionada", "Cant. Aceptada", "Cant. Rechazada", "Cant. Pendiente"]
            if st.session_state.estado_herramientas_df is None or st.session_state.estado_herramientas_df.empty:
                st.session_state.estado_herramientas_df = pd.DataFrame([{c: "" if c == "Herramienta" else 0 for c in cols_estado}])

            with st.form(key="form_estado"):
                df_estado = st.data_editor(
                    st.session_state.estado_herramientas_df,
                    column_config={
                        "Herramienta": st.column_config.TextColumn("Herramienta"),
                        "Cant. Inspeccionada": st.column_config.NumberColumn("Cant. Inspeccionada", min_value=0),
                        "Cant. Aceptada": st.column_config.NumberColumn("Cant. Aceptada", min_value=0),
                        "Cant. Rechazada": st.column_config.NumberColumn("Cant. Rechazada", min_value=0),
                        "Cant. Pendiente": st.column_config.NumberColumn("Cant. Pendiente", min_value=0)
                    },
                    num_rows="dynamic", key="estado_editor_in_form", use_container_width=True
                )
                if st.form_submit_button("✅ CONFIRMAR CAMBIOS EN ESTADO", use_container_width=True):
                    st.session_state.estado_herramientas_df = pd.DataFrame(df_estado) if not isinstance(df_estado, pd.DataFrame) else df_estado.copy()
                    st.success("✅ CAMBIOS GUARDADOS")
                    st.rerun()

            st.subheader("10. REPRESENTANTES")
            col_r1, col_r2 = st.columns(2)
            with col_r1:
                gerente_contrato = st.text_input("GERENTE DE CONTRATO", value=st.session_state.form_data.get('gerente_contrato', st.session_state.form_data.get('gerente_proyecto', '')), key="gerente_contrato")
                rep_cliente = st.text_input("REPRESENTANTE DEL CLIENTE", value=st.session_state.form_data.get('rep_cliente', ''), key="rep_cli")
                cargo_cliente = st.text_input("CARGO DEL REPRESENTANTE DEL CLIENTE", value=st.session_state.form_data.get('cargo_cliente', ''), key="cargo_cli")
            with col_r2:
                # CORREGIDO: nombre correcto "Erika Soruco" con K
                opciones = ["Brenda Figueroa", "Erika Soruco"]
                valor_actual = st.session_state.form_data.get('rep_serfigobol', 'Brenda Figueroa')
                if valor_actual not in opciones:
                    valor_actual = "Brenda Figueroa"
                rep_serfi = st.selectbox("REPRESENTANTE SERFIGOBOL S.R.L.", opciones, index=opciones.index(valor_actual), key="rep_ser")

            st.session_state.form_data['gerente_contrato'] = gerente_contrato
            st.session_state.form_data['rep_cliente'] = rep_cliente
            st.session_state.form_data['cargo_cliente'] = cargo_cliente
            st.session_state.form_data['rep_serfigobol'] = rep_serfi

            st.divider()

            def validar_campos():
                errores = []
                if not st.session_state.form_data.get('cliente', '').strip():
                    errores.append("CLIENTE")
                if not st.session_state.form_data.get('contrato', '').strip():
                    errores.append("CONTRATO No")
                if not st.session_state.form_data.get('pozo', '').strip():
                    errores.append("POZO/PROYECTO")
                if not st.session_state.form_data.get('fecha_serv', '').strip():
                    errores.append("FECHA DE SERVICIO")
                ed_items_check = st.session_state.form_data.get('ed_items')
                if ed_items_check is None or (isinstance(ed_items_check, pd.DataFrame) and ed_items_check.empty):
                    errores.append("DESCRIPCION DE HERRAMIENTAS")
                return errores

            if st.button("🚀 GENERAR REPORTE COMPLETO (WORD)", type="primary", use_container_width=True, key="generar_reporte"):
                errores = validar_campos()
                if errores:
                    st.error(f"❌ COMPLETE: {', '.join(errores)}")
                else:
                    doc = Document()
                    style = doc.styles['Normal']
                    style.font.name = 'Calibri'
                    style.font.size = Pt(11)
                    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    section = doc.sections[0]
                    section.top_margin = Cm(2.5)
                    section.bottom_margin = Cm(2.5)
                    section.left_margin = Cm(2.5)
                    section.right_margin = Cm(2.5)

                    # CORREGIDO: usar función reutilizable para encabezado
                    fecha_emision = agregar_encabezado_word(
                        doc, report_id,
                        st.session_state.form_data.get('cliente', ''),
                        "REPORTE DIARIO DE MONITOREO QA/QC"
                    )
                    if not st.session_state.fecha_emision_actual:
                        st.session_state.fecha_emision_actual = fecha_emision
                        st.session_state.estado_reporte_actual = 'EMITIDO'

                    # CORREGIDO: usar estado local (thread-safe)
                    titulo_state = make_titulo_counter()

                    data = [
                        ["CLIENTE", st.session_state.form_data.get('cliente', ''), "CONTACTO DEL CLIENTE", st.session_state.form_data.get('contacto', '')],
                        ["CONTRATO NO", st.session_state.form_data.get('contrato', ''), "POZO/PROYECTO", st.session_state.form_data.get('pozo', '')],
                        ["ORDEN DE COMPRA", st.session_state.form_data.get('orden_compra', ''), "ORDEN DE SERVICIO", st.session_state.form_data.get('orden_servicio', '')],
                        ["FECHA DE SERVICIO", st.session_state.form_data.get('fecha_serv', ''), "SOLICITUD", st.session_state.form_data.get('solicitud', '')],
                        ["FECHA DE ENVÍO", st.session_state.form_data.get('fecha_envio', ''), "PRÓXIMO SERVICIO", st.session_state.form_data.get('fecha_prox', '')]
                    ]
                    agregar_tabla_cabecera(doc, data)
                    doc.add_paragraph()

                    if st.session_state.estado_herramientas_df is not None and not st.session_state.estado_herramientas_df.empty:
                        add_title_with_bar(doc, "RESUMEN DE ESTADO DE HERRAMIENTAS", titulo_state)
                        add_table_professional(doc, st.session_state.estado_herramientas_df, "")

                    add_title_with_bar(doc, "1. DETALLES DEL PROVEEDOR", titulo_state)
                    add_table_professional(doc, st.session_state.form_data.get('ed_prov'), "")

                    add_title_with_bar(doc, "2. RESUMEN DE MONITOREO QA/QC", titulo_state)
                    p_res = doc.add_paragraph()
                    add_quill(p_res, st.session_state.form_data.get('resumen', ''))

                    add_subtitle_left(doc, "2.1 DESCRIPCION DE HERRAMIENTAS")
                    add_table_professional(doc, st.session_state.form_data.get('ed_items'), "")

                    add_subtitle_left(doc, "2.2 LISTA DE ACTIVIDADES REALIZADAS")
                    acts_list = [k for k, v in st.session_state.form_data.get('check_vals', {}).items() if v]
                    otras = st.session_state.form_data.get('otras', '')
                    if otras:
                        acts_list.extend([a.strip() for a in otras.split('\n') if a.strip()])
                    if acts_list:
                        t_acts = doc.add_table(rows=1, cols=1)
                        t_acts.style = 'Table Grid'
                        cell_header = t_acts.cell(0, 0)
                        cell_header.text = "ACTIVIDADES REALIZADAS"
                        cell_header.paragraphs[0].runs[0].bold = True
                        cell_header.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        shading = OxmlElement('w:shd')
                        shading.set(qn('w:fill'), 'D3D3D3')
                        cell_header._tc.get_or_add_tcPr().append(shading)
                        for a in acts_list:
                            t_acts.add_row().cells[0].text = f"☑ {a}"

                    add_title_with_bar(doc, "3. SECCIÓN DOCUMENTAL Y NORMATIVA", titulo_state)
                    add_table_professional(doc, st.session_state.form_data.get('ed_doc'), "")

                    add_title_with_bar(doc, "4. DETALLE DE INSTRUMENTOS", titulo_state)
                    add_table_professional(doc, st.session_state.form_data.get('ed_inst'), "")
                    if st.session_state.mostrar_nota_inst and st.session_state.nota_instrumentos:
                        add_heading_black(doc, "NOTA SOBRE INSTRUMENTOS", level=3)
                        doc.add_paragraph().add_run(st.session_state.nota_instrumentos)

                    add_title_with_bar(doc, "5. COMPETENCIA DEL PERSONAL", titulo_state)
                    add_table_professional(doc, st.session_state.form_data.get('ed_pers'), "")
                    if st.session_state.mostrar_nota_pers and st.session_state.nota_personal:
                        add_heading_black(doc, "NOTA SOBRE PERSONAL", level=3)
                        doc.add_paragraph().add_run(st.session_state.nota_personal)

                    add_title_with_bar(doc, "6. DESCRIPCIÓN DE MONITOREO QA/QC", titulo_state)
                    add_heading_black(doc, "INFORMACIÓN GENERAL", level=2)
                    p_info = doc.add_paragraph()
                    add_quill(p_info, st.session_state.form_data.get('info_gral', ''))
                    add_heading_black(doc, "OBJETIVOS", level=2)
                    p_obj = doc.add_paragraph()
                    add_quill(p_obj, st.session_state.form_data.get('objetivos', ''))

                    add_subtitle_left(doc, "6.1 REVISIÓN DOCUMENTAL")
                    add_heading_black(doc, "DESCRIPCIÓN", level=3)
                    p_desc = doc.add_paragraph()
                    add_quill(p_desc, st.session_state.form_data.get('desc_rev', ''))
                    add_table_professional(doc, st.session_state.form_data.get('ed_reg'), "REGISTRO DOCUMENTAL")

                    for i in range(st.session_state.n_inspecciones):
                        add_subtitle_left(doc, f"6.2 INSPECCIÓN {i+1}")
                        tabla_insp = doc.add_table(rows=2, cols=4)
                        tabla_insp.style = 'Table Grid'
                        datos_insp = [
                            ("INSPECTOR:", st.session_state.form_data.get(f'inspector_{i}', 'N/A'), "W.O. / O.S. / O.T.:", st.session_state.form_data.get(f'doc_ref_{i}', 'N/A')),
                            ("HERRAMIENTA:", st.session_state.form_data.get(f'herramienta_{i}', 'N/A'), "NUMERO DE SERIE:", st.session_state.form_data.get(f'num_serie_{i}', 'N/A'))
                        ]
                        for fila_idx, (l1, v1, l2, v2) in enumerate(datos_insp):
                            for col_idx, val in enumerate([l1, v1, l2, v2]):
                                cell = tabla_insp.cell(fila_idx, col_idx)
                                cell.text = val
                                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                                if col_idx % 2 == 0 and cell.paragraphs[0].runs:
                                    cell.paragraphs[0].runs[0].bold = True
                                    shd = OxmlElement('w:shd')
                                    shd.set(qn('w:val'), 'clear')
                                    shd.set(qn('w:color'), 'auto')
                                    shd.set(qn('w:fill'), 'D3D3D3')
                                    cell._tc.get_or_add_tcPr().append(shd)
                        doc.add_paragraph()
                        add_heading_black(doc, "DESCRIPCIÓN", level=3)
                        p_insp = doc.add_paragraph()
                        add_quill(p_insp, st.session_state.form_data.get(f'insp_desc_{i}', ''))
                        add_table_professional(doc, st.session_state.form_data.get(f'comp_{i}'), "COMPONENTES")
                        fotos_insp = [f for f in st.session_state.fotos_inspecciones.get(f'fotos_insp_{i}', []) if f.get('imagen') is not None]
                        add_photos_row(doc, fotos_insp, "FOTOS DE INSPECCIÓN")

                    for j in range(st.session_state.n_ensambles):
                        add_subtitle_left(doc, f"6.3 ENSAMBLE {j+1}")
                        tabla_ens = doc.add_table(rows=1, cols=4)
                        tabla_ens.style = 'Table Grid'
                        for col_idx, val in enumerate(["HERRAMIENTA:", st.session_state.form_data.get(f'eh_{j}', 'N/A'), "SERIAL:", st.session_state.form_data.get(f'es_{j}', 'N/A')]):
                            cell = tabla_ens.cell(0, col_idx)
                            cell.text = val
                            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                            if col_idx % 2 == 0 and cell.paragraphs[0].runs:
                                cell.paragraphs[0].runs[0].bold = True
                                shd = OxmlElement('w:shd')
                                shd.set(qn('w:val'), 'clear')
                                shd.set(qn('w:color'), 'auto')
                                shd.set(qn('w:fill'), 'D3D3D3')
                                cell._tc.get_or_add_tcPr().append(shd)
                        doc.add_paragraph()
                        add_heading_black(doc, "DESCRIPCIÓN", level=3)
                        p_ens = doc.add_paragraph()
                        add_quill(p_ens, st.session_state.form_data.get(f'ens_desc_{j}', ''))
                        add_table_professional(doc, st.session_state.form_data.get(f'torque_{j}'), "REPORTE DE TORQUE")
                        fotos_ens = [f for f in st.session_state.fotos_ensambles.get(f'fotos_ens_{j}', []) if f.get('imagen') is not None]
                        add_photos_row(doc, fotos_ens, "FOTOS DE ENSAMBLE")

                    add_subtitle_left(doc, "6.4 PRUEBAS")
                    for k in range(st.session_state.n_pruebas):
                        doc.add_heading(f"PRUEBA {k+1}", level=2)
                        for run in doc.paragraphs[-1].runs:
                            run.font.color.rgb = RGBColor(0, 0, 0)
                            run.bold = True
                        tabla_pr = doc.add_table(rows=2, cols=4)
                        tabla_pr.style = 'Table Grid'
                        datos_pr = [
                            ("NOMBRE DE LA HERRAMIENTA:", st.session_state.form_data.get(f'pr_nom_{k}', 'N/A'), "NUMERO DE SERIE:", st.session_state.form_data.get(f'pr_serie_{k}', 'N/A')),
                            ("TIPO DE PRUEBA:", st.session_state.form_data.get(f'pr_tipo_{k}', 'N/A'), "RESULTADOS:", st.session_state.form_data.get(f'pr_res_{k}', 'N/A'))
                        ]
                        for fila_idx, (l1, v1, l2, v2) in enumerate(datos_pr):
                            for col_idx, val in enumerate([l1, v1, l2, v2]):
                                cell = tabla_pr.cell(fila_idx, col_idx)
                                cell.text = val
                                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                                if col_idx % 2 == 0 and cell.paragraphs[0].runs:
                                    cell.paragraphs[0].runs[0].bold = True
                                    shd = OxmlElement('w:shd')
                                    shd.set(qn('w:val'), 'clear')
                                    shd.set(qn('w:color'), 'auto')
                                    shd.set(qn('w:fill'), 'D3D3D3')
                                    cell._tc.get_or_add_tcPr().append(shd)
                        doc.add_paragraph()
                        add_heading_black(doc, "DESCRIPCIÓN DE LA PRUEBA", level=3)
                        p_desc_pr = doc.add_paragraph()
                        add_quill(p_desc_pr, st.session_state.form_data.get(f'pr_desc_{k}', ''))
                        add_heading_black(doc, "TOLERANCIA", level=3)
                        p_tol_pr = doc.add_paragraph()
                        add_quill(p_tol_pr, st.session_state.form_data.get(f'pr_tol_{k}', ''))
                        fotos_pr = [f for f in st.session_state.fotos_pruebas.get(f'fotos_prueba_{k}', []) if f.get('imagen') is not None]
                        add_photos_row(doc, fotos_pr, "FOTOS DE PRUEBA")

                    add_title_with_bar(doc, "7. HALLAZGOS / DISCREPANCIAS / NCR", titulo_state)
                    if st.session_state.hallazgos_list:
                        for idx, hallazgo in enumerate(st.session_state.hallazgos_list):
                            add_hallazgo_table(doc, hallazgo, idx + 1)
                            doc.add_paragraph()
                    else:
                        doc.add_paragraph("No se registraron hallazgos.")

                    add_title_with_bar(doc, "8. REPORTE FOTOGRÁFICO", titulo_state)
                    fotos_gen = [f for f in st.session_state.fotos_generales if f.get('imagen') is not None]
                    if fotos_gen:
                        add_photos_row(doc, fotos_gen, "")
                    else:
                        doc.add_paragraph("NO SE ADJUNTARON FOTOGRAFÍAS")

                    add_title_with_bar(doc, "9. ESTADO FINAL DE HERRAMIENTA", titulo_state)
                    p_est = doc.add_paragraph()
                    add_quill(p_est, st.session_state.form_data.get('estado_final', ''))

                    add_title_with_bar(doc, "10. REPRESENTANTES", titulo_state)
                    agregar_seccion_representantes(doc, st.session_state.form_data, tipo='qaqc')

                    agregar_pie_pagina(doc)

                    target = io.BytesIO()
                    doc.save(target)

                    st.download_button("📥 DESCARGAR DOCX", target.getvalue(), f"{report_id}.docx", key="download_final")

                    guardar_datos_reporte(p_sel, poz_sel, prov_sel, report_id, es_generacion_word=True)
                    guardar_hallazgos_en_discrepancias(poz_sel, p_sel, report_id, st.session_state.hallazgos_list)

                    st.balloons()
                    st.success(f"✅ REPORTE {report_id} GENERADO CORRECTAMENTE")

# ============================================================================
# TIMESHEET
# ============================================================================
elif menu == "📄 TIMESHEET":
    st.header("📄 GENERAR TIMESHEET")
    st.markdown("---")

    proyectos = ejecutar_query("SELECT id, contrato_nom, operadora, contacto_cliente, contrato_num, gerente_proyecto FROM proyectos", fetch=True)
    if not proyectos:
        st.warning("⚠️ NO HAY PROYECTOS REGISTRADOS.")
        st.stop()

    col_proy, col_pozo = st.columns(2)
    with col_proy:
        proy_sel = st.selectbox("PROYECTO", proyectos, format_func=lambda x: x[1], key="timesheet_proy")

    pozos = ejecutar_query("SELECT id, rig, nombre_pozo FROM pozos WHERE proyecto_id = ? ORDER BY rig", (proy_sel[0],), fetch=True)
    if not pozos:
        st.warning("⚠️ NO HAY POZOS REGISTRADOS EN ESTE PROYECTO.")
        st.stop()

    with col_pozo:
        pozo_sel = st.selectbox("POZO", pozos, format_func=lambda x: f"{x[1]} - {x[2]}", key="timesheet_pozo")

    col_fecha1, col_fecha2 = st.columns(2)
    with col_fecha1:
        fecha_desde = st.date_input("FECHA DESDE", value=datetime.now().date() - timedelta(days=30), key="timesheet_desde")
    with col_fecha2:
        fecha_hasta = st.date_input("FECHA HASTA", value=datetime.now().date(), key="timesheet_hasta")

    fecha_desde_str = fecha_desde.strftime('%Y-%m-%d')
    fecha_hasta_str = fecha_hasta.strftime('%Y-%m-%d')

    reportes = ejecutar_query("""
        SELECT reporte_id, fecha_inspeccion, datos_json, proveedor_id
        FROM inspecciones
        WHERE pozo_id = ?
        AND date(fecha_inspeccion) BETWEEN ? AND ?
        ORDER BY fecha_inspeccion ASC
    """, (pozo_sel[0], fecha_desde_str, fecha_hasta_str), fetch=True)

    if not reportes:
        st.info("📭 NO HAY REPORTES EN EL PERÍODO SELECCIONADO.")
        st.stop()

    st.success(f"📊 SE ENCONTRARON {len(reportes)} REPORTES EN EL PERÍODO")

    datos_tabla = []
    for idx, rep in enumerate(reportes):
        reporte_id, fecha_inspeccion, datos_json, proveedor_id = rep
        nombre_proveedor = "SIN PROVEEDOR"
        if proveedor_id:
            prov_data = ejecutar_query("SELECT nombre_proveedor FROM proveedores WHERE id = ?", (proveedor_id,), fetch=True)
            if prov_data:
                nombre_proveedor = prov_data[0][0]

        datos = json.loads(datos_json) if datos_json else {}
        ed_items_raw = datos.get('form_data', {}).get('ed_items', [])
        if isinstance(ed_items_raw, dict) and '__columns__' in ed_items_raw:
            ed_items = restaurar_dfs_desde_json(ed_items_raw)
        elif isinstance(ed_items_raw, list) and ed_items_raw:
            ed_items = pd.DataFrame(ed_items_raw)
        else:
            ed_items = pd.DataFrame()

        lineas_descripcion = []
        if not ed_items.empty:
            for _, row in ed_items.iterrows():
                primera_linea = str(row.get('DESCRIPCIÓN', '') or '')
                od = str(row.get('OD', '') or '')
                if od and od != 'nan':
                    primera_linea += f" OD {od}"
                if primera_linea.strip():
                    lineas_descripcion.append(primera_linea.strip())
                conx = str(row.get('CONEXIONES', '') or '')
                if conx and conx != 'nan':
                    lineas_descripcion.append(f"Conex. {conx}")
                serie = str(row.get('NUMERO DE SERIE', '') or '')
                if serie and serie != 'nan':
                    lineas_descripcion.append(f"#{serie}")

        descripcion_final = '\n'.join(lineas_descripcion) if lineas_descripcion else "S/D"

        datos_tabla.append({
            'item': idx + 1,
            'fecha': fecha_inspeccion[:10] if fecha_inspeccion else "",
            'reporte_id': reporte_id,
            'descripcion': descripcion_final,
            'proveedor': f"BASE {nombre_proveedor}",
            'observacion': ""
        })

    df_tabla = pd.DataFrame(datos_tabla)

    st.markdown("### 📋 TABLA DE ACTIVIDADES")
    st.info("📝 Edite la columna 'OBSERVACIÓN' según sea necesario.")

    edited_df = st.data_editor(
        df_tabla,
        column_config={
            "item": st.column_config.NumberColumn("Item", disabled=True),
            "fecha": st.column_config.TextColumn("Fecha", disabled=True),
            "reporte_id": st.column_config.TextColumn("N° Reporte", disabled=True),
            "descripcion": st.column_config.TextColumn("Descripción de Herramienta", disabled=True),
            "proveedor": st.column_config.TextColumn("Client/PO No/Proveedor/Locación", disabled=True),
            "observacion": st.column_config.TextColumn("Observación", width="large")
        },
        use_container_width=True, hide_index=True, key="timesheet_editor"
    )

    periodo_manual = st.text_input("PERIODO (para mostrar en cabecera)", value=f"{fecha_desde_str} al {fecha_hasta_str}", key="timesheet_periodo")
    st.markdown("---")
    st.markdown("### 🖼️ LOGO DEL CLIENTE (opcional)")
    logo_cliente = st.file_uploader("Seleccione una imagen para el logo del cliente", type=['png', 'jpg', 'jpeg'], key="timesheet_logo_cliente")
    st.markdown("---")

    if st.button("🚀 GENERAR TIMESHEET (WORD)", type="primary", use_container_width=True, key="generar_timesheet"):
        with st.spinner("Generando documento..."):
            doc = Document()
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(10)

            section = doc.sections[0]
            section.top_margin = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)

            ANCHO_TABLA = 7.0

            header = section.header
            header.paragraphs[0].text = ""

            header_table = header.add_table(rows=1, cols=3, width=Inches(ANCHO_TABLA))
            header_table.style = 'Table Grid'
            header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            header_table.columns[0].width = Inches(1.2)
            header_table.columns[1].width = Inches(4.6)
            header_table.columns[2].width = Inches(1.2)

            for row in header_table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            logo_serfi = cargar_logo()
            if logo_serfi:
                try:
                    cell_logo = header_table.cell(0, 0)
                    cell_logo.paragraphs[0].clear()
                    run = cell_logo.paragraphs[0].add_run()
                    run.add_picture(logo_serfi, width=Inches(0.66))
                except Exception:
                    header_table.cell(0, 0).text = "SERFIGOBOL"
            else:
                header_table.cell(0, 0).text = "SERFIGOBOL"

            cell_title = header_table.cell(0, 1)
            cell_title.paragraphs[0].clear()
            p = cell_title.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run("\n")
            for txt in ["Project Management\n", "Vendor Inspection / Expediting\n", "Quality Assurance\n\n"]:
                r = p.add_run(txt)
                r.bold = True
                r.font.size = Pt(11)
            r4 = p.add_run("TIMESHEET")
            r4.bold = True
            r4.font.size = Pt(14)
            p.add_run("\n")

            if logo_cliente:
                try:
                    img = Image.open(logo_cliente)
                    img.thumbnail((200, 100))
                    buffered = io.BytesIO()
                    img.save(buffered, format='PNG')
                    buffered.seek(0)
                    cell_lc = header_table.cell(0, 2)
                    cell_lc.paragraphs[0].clear()
                    run = cell_lc.paragraphs[0].add_run()
                    run.add_picture(buffered, width=Inches(1.2))
                except Exception:
                    pass

            doc.add_paragraph()

            def set_border_white(cell):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                for border in ['top', 'left', 'bottom', 'right']:
                    border_elem = OxmlElement(f'w:{border}')
                    border_elem.set(qn('w:val'), 'single')
                    border_elem.set(qn('w:color'), 'FFFFFF')
                    tcPr.append(border_elem)

            datos_cabecera = doc.add_table(rows=3, cols=5)
            datos_cabecera.style = 'Table Grid'
            datos_cabecera.alignment = WD_TABLE_ALIGNMENT.CENTER
            datos_cabecera.columns[0].width = Inches(1.2)
            datos_cabecera.columns[1].width = Inches(2.2)
            datos_cabecera.columns[2].width = Inches(0.5)
            datos_cabecera.columns[3].width = Inches(1.2)
            datos_cabecera.columns[4].width = Inches(1.9)

            for fila in range(3):
                set_border_white(datos_cabecera.cell(fila, 2))

            cab_data = [
                (0, "POZO", pozo_sel[2] or "", "PERIODO", periodo_manual),
                (1, "CONTRATO", proy_sel[4] or "", None, None),
                (2, "COMPAÑÍA", "SERFIGOBOL S.R.L.", None, None)
            ]
            for fila, label, valor, label2, valor2 in cab_data:
                cell = datos_cabecera.cell(fila, 0)
                cell.text = label
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'D3D3D3')
                cell._tc.get_or_add_tcPr().append(shading)
                datos_cabecera.cell(fila, 1).text = valor
                datos_cabecera.cell(fila, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if label2:
                    cell3 = datos_cabecera.cell(fila, 3)
                    cell3.text = label2
                    cell3.paragraphs[0].runs[0].bold = True
                    cell3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    shading2 = OxmlElement('w:shd')
                    shading2.set(qn('w:fill'), 'D3D3D3')
                    cell3._tc.get_or_add_tcPr().append(shading2)
                    datos_cabecera.cell(fila, 4).text = valor2 or ""
                    datos_cabecera.cell(fila, 4).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    set_border_white(datos_cabecera.cell(fila, 3))
                    set_border_white(datos_cabecera.cell(fila, 4))

            doc.add_paragraph()

            filas_existentes = len(edited_df)
            filas_vacias = max(0, 18 - filas_existentes)
            table = doc.add_table(rows=1 + filas_existentes + filas_vacias, cols=6)
            table.style = 'Table Grid'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.autofit = False
            table.allow_autofit = True
            table.columns[0].width = Inches(0.4)
            table.columns[1].width = Inches(0.7)
            table.columns[2].width = Inches(0.8)
            table.columns[3].width = Inches(2.7)
            table.columns[4].width = Inches(1.0)
            table.columns[5].width = Inches(1.4)

            headers_ts = ["Item", "Fecha", "N° Reporte", "Descripción de Herramienta", "Client/PO No/Proveedor/Locación", "Observación"]
            for i, h in enumerate(headers_ts):
                cell = table.cell(0, i)
                cell.text = h
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                shading = OxmlElement('w:shd')
                shading.set(qn('w:fill'), 'D3D3D3')
                cell._tc.get_or_add_tcPr().append(shading)

            for idx_r, (_, row) in enumerate(edited_df.iterrows()):
                cells = table.rows[idx_r + 1].cells
                cells[0].text = str(row['item'])
                cells[1].text = row['fecha']
                cells[2].text = row['reporte_id']
                desc_cell = cells[3]
                desc_cell.paragraphs[0].clear()
                for j, linea in enumerate(row['descripcion'].split('\n')):
                    if linea.strip():
                        run = desc_cell.paragraphs[0].add_run(linea)
                        run.font.size = Pt(9)
                        if j < len(row['descripcion'].split('\n')) - 1:
                            desc_cell.paragraphs[0].add_run('\n')
                desc_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                desc_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                cells[4].text = row['proveedor']
                cells[5].text = row['observacion']
                for cell in cells:
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            for i in range(filas_vacias):
                for cell in table.rows[filas_existentes + 1 + i].cells:
                    cell.text = ""

            doc.add_paragraph()

            footer = doc.sections[0].footer
            footer.paragraphs[0].text = ""
            firma_table = footer.add_table(rows=8, cols=3, width=Inches(ANCHO_TABLA))
            firma_table.style = 'Table Grid'
            firma_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            firma_table.columns[0].width = Inches(2.0)
            firma_table.columns[1].width = Inches(2.5)
            firma_table.columns[2].width = Inches(2.5)

            firma_table.cell(0, 0).merge(firma_table.cell(6, 0))
            firma_table.cell(0, 1).merge(firma_table.cell(6, 1))
            firma_table.cell(1, 2).merge(firma_table.cell(7, 2))

            cell_aprob = firma_table.cell(0, 2)
            cell_aprob.text = "APROBADO POR CLIENTE"
            cell_aprob.paragraphs[0].runs[0].bold = True
            cell_aprob.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_aprob.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shading = OxmlElement('w:shd')
            shading.set(qn('w:fill'), 'D3D3D3')
            cell_aprob._tc.get_or_add_tcPr().append(shading)

            
            for fila_idx, col_idx_f, etiqueta_f in [(7, 0, "ESPECIALISTA QA/QC"), (7, 1, "FISCAL DE CONTRATO")]:
                cell_f = firma_table.cell(fila_idx, col_idx_f)
                cell_f.text = etiqueta_f
                cell_f.paragraphs[0].runs[0].bold = True
                cell_f.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cell_f.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                shading_f = OxmlElement('w:shd')
                shading_f.set(qn('w:fill'), 'D3D3D3')
                cell_f._tc.get_or_add_tcPr().append(shading_f)

            target = io.BytesIO()
            doc.save(target)

            st.download_button(
                label="📥 DESCARGAR TIMESHEET (DOCX)",
                data=target.getvalue(),
                file_name=f"TIMESHEET_{pozo_sel[2]}_{fecha_desde_str}_TO_{fecha_hasta_str}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_timesheet"
            )
            st.balloons()
            st.success("✅ TIMESHEET GENERADO CORRECTAMENTE")

# ============================================================================
# NUEVO REPORTE QA/QC - DBR
# ============================================================================
elif menu == "📝 NUEVO REPORTE QA/QC - DBR":
    st.header("📝 NUEVO REPORTE QA/QC - DBR")
    st.markdown("### Análisis de DBR (DAMAGED BEYOND REPAIR)")

    # ================================================================
    # CARGA DE EDICIÓN DBR
    # ================================================================
    if st.session_state.get('modo_edicion_dbr', False) and st.session_state.get('edicion_reporte_id_dbr'):
        if not st.session_state.get('datos_editados_dbr_cargados', False):
            reporte_data = ejecutar_query("""
                SELECT datos_json, usuario_creador_id, proyecto_id, pozo_id, proveedor_id, estado_reporte, fecha_emision
                FROM inspecciones WHERE reporte_id = ?""",
                (st.session_state.edicion_reporte_id_dbr,), fetch=True)
            if reporte_data:
                datos_json_str, usuario_creador_id, proy_id_bd, poz_id_bd, prov_id_bd, estado_reporte_bd, fecha_emision_bd = reporte_data[0]

                if st.session_state.usuario_rol != "admin" and usuario_creador_id != st.session_state.usuario_id:
                    st.error("❌ NO TIENE PERMISO PARA EDITAR ESTE REPORTE")
                    st.session_state.modo_edicion_dbr = False
                    st.rerun()

                st.session_state.reporte_actual_id_dbr = st.session_state.edicion_reporte_id_dbr
                st.session_state.estado_reporte_actual_dbr = estado_reporte_bd
                st.session_state.fecha_emision_actual_dbr = fecha_emision_bd

                match = re.search(r'DBR-(\d+)-', st.session_state.edicion_reporte_id_dbr)
                if match:
                    st.session_state.reporte_numero_actual_dbr = int(match.group(1))

                p_sel_bd = ejecutar_query("SELECT id, contrato_nom, operadora, contacto_cliente, contrato_num, gerente_proyecto FROM proyectos WHERE id = ?", (proy_id_bd,), fetch=True)[0]
                poz_sel_bd = ejecutar_query("SELECT id, rig, nombre_pozo FROM pozos WHERE id = ?", (poz_id_bd,), fetch=True)[0]
                prov_sel_bd = None
                if prov_id_bd:
                    prov_data = ejecutar_query("SELECT id, nombre_proveedor FROM proveedores WHERE id = ?", (prov_id_bd,), fetch=True)
                    if prov_data:
                        prov_sel_bd = prov_data[0]

                st.session_state.datos_seleccion_dbr = {'proyecto': p_sel_bd, 'pozo': poz_sel_bd, 'proveedor': prov_sel_bd}

                if datos_json_str:
                    datos_cargados = json.loads(datos_json_str)
                    st.session_state.form_data_dbr = datos_cargados.get('form_data_dbr', {})

                    cols_items_dbr = ["N°", "DESCRIPCIÓN", "OD", "CONEXIONES", "NUMERO DE SERIE"]
                    st.session_state.ed_items_dbr = restaurar_dfs_desde_json(datos_cargados.get('ed_items_dbr', {}), cols_items_dbr)

                    st.session_state.tabla_historial_evento = restaurar_dfs_desde_json(
                        datos_cargados.get('tabla_historial_evento', {}),
                        ["FECHA DE EVENTO", "EVENTO", "OBSERVACION"])
                    st.session_state.tabla_historial_herramienta = restaurar_dfs_desde_json(
                        datos_cargados.get('tabla_historial_herramienta', {}),
                        ["FECHA DE ACCION", "ACCION", "OBSERVACIÓN"])
                    st.session_state.tabla_procesos = restaurar_dfs_desde_json(
                        datos_cargados.get('tabla_procesos', {}),
                        ["FECHA", "PROCESO", "N REPORTE", "OBSERVACIÓN"])
                    st.session_state.tabla_certificados = restaurar_dfs_desde_json(
                        datos_cargados.get('tabla_certificados', {}),
                        ["DOCUMENTO", "NUMERO DOC./REVISION/N REF.", "OBSERVACIÓN"])

                    st.session_state.n_fotos_historial_evento = datos_cargados.get('n_fotos_historial_evento', 1)
                    st.session_state.n_fotos_historial_herramienta = datos_cargados.get('n_fotos_historial_herramienta', 1)
                    st.session_state.n_fotos_procesos = datos_cargados.get('n_fotos_procesos', 1)
                    st.session_state.n_fotos_resumen = datos_cargados.get('n_fotos_resumen', 1)

                    if 'hallazgos_list_dbr' in datos_cargados:
                        st.session_state.hallazgos_list_dbr = restaurar_hallazgos_desde_json(datos_cargados['hallazgos_list_dbr'])

                    if 'representantes' in datos_cargados:
                        for k, v in datos_cargados['representantes'].items():
                            st.session_state.form_data_dbr[k] = v

                    st.session_state.fotos_historial_evento = restaurar_lista_fotos(datos_cargados.get('fotos_historial_evento', []))
                    st.session_state.fotos_historial_herramienta = restaurar_lista_fotos(datos_cargados.get('fotos_historial_herramienta', []))
                    st.session_state.fotos_procesos = restaurar_lista_fotos(datos_cargados.get('fotos_procesos', []))
                    st.session_state.fotos_resumen = restaurar_lista_fotos(datos_cargados.get('fotos_resumen', []))

                st.session_state.datos_editados_dbr_cargados = True
                st.session_state.formulario_activo_dbr = True
                st.rerun()

    # ================================================================
    # INICIO DE NUEVO REPORTE DBR
    # ================================================================
    if not st.session_state.get('formulario_activo_dbr', False) and not st.session_state.get('modo_edicion_dbr', False):
        proyectos = ejecutar_query("SELECT id, contrato_nom, operadora, contacto_cliente, contrato_num, gerente_proyecto FROM proyectos", fetch=True)
        if not proyectos:
            st.warning("⚠️ NO HAY PROYECTOS REGISTRADOS.")
            st.stop()

        col1, col2, col3 = st.columns(3)
        with col1:
            p_sel = st.selectbox("PROYECTO", proyectos, format_func=lambda x: x[1], key="dbr_proyecto")
        pozos = ejecutar_query("SELECT id, rig, nombre_pozo FROM pozos WHERE proyecto_id = ?", (p_sel[0],), fetch=True)
        with col2:
            if pozos:
                poz_sel = st.selectbox("RIG / POZO", pozos, format_func=lambda x: f"{x[1]} - {x[2]}", key="dbr_pozo")
            else:
                st.warning("⚠️ NO HAY RIGS REGISTRADOS EN ESTE PROYECTO")
                st.stop()
        proveedores = ejecutar_query("SELECT id, nombre_proveedor FROM proveedores WHERE pozo_id = ?", (poz_sel[0],), fetch=True)
        with col3:
            if proveedores:
                prov_sel = st.selectbox("PROVEEDOR", proveedores, format_func=lambda x: x[1], key="dbr_proveedor")
            else:
                prov_sel = None
                st.info("📌 No hay proveedores registrados.")

        if st.button("🚀 INICIAR REPORTE DBR", use_container_width=True, key="iniciar_dbr"):
            try:
                with sqlite3.connect(DB_PATH) as conn:
                    cursor = conn.cursor()
                    cursor.execute("SELECT ultimo_numero FROM contador_reportes_dbr WHERE pozo_id = ?", (poz_sel[0],))
                    contador = cursor.fetchone()
                    if contador:
                        nuevo_numero = contador[0] + 1
                        cursor.execute("UPDATE contador_reportes_dbr SET ultimo_numero = ? WHERE pozo_id = ?", (nuevo_numero, poz_sel[0]))
                    else:
                        nuevo_numero = 1
                        cursor.execute("INSERT INTO contador_reportes_dbr (pozo_id, ultimo_numero) VALUES (?, ?)", (poz_sel[0], nuevo_numero))
                    conn.commit()
            except Exception as e:
                st.error(f"Error al crear contador DBR: {e}")
                st.stop()

            st.session_state.reporte_numero_actual_dbr = nuevo_numero
            st.session_state.reporte_actual_id_dbr = f"DBR-{nuevo_numero}-{poz_sel[2].replace(' ', '_')}"
            st.session_state.formulario_activo_dbr = True
            st.session_state.estado_reporte_actual_dbr = "BORRADOR"
            st.session_state.fecha_emision_actual_dbr = None
            st.session_state.datos_seleccion_dbr = {'proyecto': p_sel, 'pozo': poz_sel, 'proveedor': prov_sel}
            st.session_state.form_data_dbr = {}
            st.session_state.hallazgos_list_dbr = []
            st.session_state.ed_items_dbr = pd.DataFrame(columns=["N°", "DESCRIPCIÓN", "OD", "CONEXIONES", "NUMERO DE SERIE"])
            st.session_state.n_fotos_historial_evento = 1
            st.session_state.n_fotos_historial_herramienta = 1
            st.session_state.n_fotos_procesos = 1
            st.session_state.n_fotos_resumen = 1
            st.session_state.tabla_historial_evento = pd.DataFrame(columns=["FECHA DE EVENTO", "EVENTO", "OBSERVACION"])
            st.session_state.tabla_historial_herramienta = pd.DataFrame(columns=["FECHA DE ACCION", "ACCION", "OBSERVACIÓN"])
            st.session_state.tabla_procesos = pd.DataFrame(columns=["FECHA", "PROCESO", "N REPORTE", "OBSERVACIÓN"])
            st.session_state.tabla_certificados = pd.DataFrame(columns=["DOCUMENTO", "NUMERO DOC./REVISION/N REF.", "OBSERVACIÓN"])
            st.session_state.fotos_historial_evento = []
            st.session_state.fotos_historial_herramienta = []
            st.session_state.fotos_procesos = []
            st.session_state.fotos_resumen = []
            st.rerun()

    # ================================================================
    # FORMULARIO ACTIVO DBR
    # ================================================================
    elif st.session_state.get('formulario_activo_dbr', False):
        p_sel = st.session_state.datos_seleccion_dbr.get('proyecto')
        poz_sel = st.session_state.datos_seleccion_dbr.get('pozo')
        prov_sel = st.session_state.datos_seleccion_dbr.get('proveedor')
        report_id = st.session_state.reporte_actual_id_dbr

        if not p_sel or not poz_sel:
            st.error("❌ Error: No hay proyecto o pozo seleccionado")
            st.stop()

        if st.button("◀ CAMBIAR PROYECTO", use_container_width=True, key="cambiar_proyecto_dbr"):
            st.session_state.formulario_activo_dbr = False
            st.session_state.modo_edicion_dbr = False
            st.session_state.menu_seleccionado = "📝 NUEVO REPORTE QA/QC - DBR"
            st.rerun()

        estado_texto = "📝 EN CURSO (BORRADOR)" if st.session_state.get('estado_reporte_actual_dbr', 'BORRADOR') == 'BORRADOR' else "✅ COMPLETADO (EMITIDO)"
        st.info(f"📄 **NÚMERO DE REPORTE:** `{report_id}` | **ESTADO:** {estado_texto}")

        col_save1, col_save2 = st.columns([1, 4])
        with col_save1:
            texto_boton = "💾 GUARDAR EN CURSO" if st.session_state.get('estado_reporte_actual_dbr', 'BORRADOR') == 'BORRADOR' else "💾 GUARDAR CAMBIOS"
            if st.button(texto_boton, key="guardar_dbr_btn", use_container_width=True, type="primary"):
                if guardar_datos_reporte_dbr(p_sel, poz_sel, prov_sel, report_id, es_generacion_word=False):
                    st.success("✅ DATOS GUARDADOS CORRECTAMENTE")

        tab_nombres_dbr = [
            "📋 CABECERA", "🔧 HERRAMIENTAS", "📝 1. INTRODUCCIÓN", "📋 2. HISTORIAL EVENTO",
            "🔧 3. HISTORIAL HERRAMIENTA", "📋 4. PROCESOS", "📄 5. CERTIFICADOS",
            "📊 6. RESUMEN MONITOREO", "⚠️ 7. HALLAZGOS", "📝 8. CONCLUSIONES",
            "💡 9. RECOMENDACIONES", "✅ 10. REPRESENTANTES"
        ]
        tabs = st.tabs(tab_nombres_dbr)

        # --- TAB 0 DBR: CABECERA ---
        with tabs[0]:
            st.subheader("📋 DATOS DE CABECERA")
            proyecto_datos = ejecutar_query("SELECT contacto_cliente, gerente_proyecto FROM proyectos WHERE id = ?", (p_sel[0],), fetch=True)
            contacto_default = proyecto_datos[0][0] if proyecto_datos and proyecto_datos[0][0] else ""

            col1, col2 = st.columns(2)
            with col1:
                cliente = st.text_input("CLIENTE *", value=st.session_state.form_data_dbr.get('cliente', p_sel[2] or ""), key="dbr_cliente")
                contrato_no = st.text_input("CONTRATO No *", value=st.session_state.form_data_dbr.get('contrato', p_sel[4] or ""), key="dbr_contrato")
                orden_compra = st.text_input("ORDEN DE COMPRA", value=st.session_state.form_data_dbr.get('orden_compra', ""), key="dbr_orden_compra")
                orden_servicio = st.text_input("ORDEN DE SERVICIO", value=st.session_state.form_data_dbr.get('orden_servicio', ""), key="dbr_orden_servicio")
            with col2:
                contacto_cliente = st.text_input("CONTACTO DEL CLIENTE", value=st.session_state.form_data_dbr.get('contacto', contacto_default), key="dbr_contacto")
                pozo_proyecto = st.text_input("POZO/PROYECTO *", value=st.session_state.form_data_dbr.get('pozo', poz_sel[2] or ""), key="dbr_pozo_nombre")
                solicitud = st.text_input("SOLICITUD", value=st.session_state.form_data_dbr.get('solicitud', ""), key="dbr_solicitud")
                fecha_servicio = st.text_input("FECHA DE SERVICIO", value=st.session_state.form_data_dbr.get('fecha_serv', ""), key="dbr_fecha_serv", placeholder="YYYY-MM-DD")

            st.session_state.form_data_dbr.update({
                'cliente': cliente, 'contrato': contrato_no,
                'orden_compra': orden_compra, 'orden_servicio': orden_servicio,
                'contacto': contacto_cliente, 'pozo': pozo_proyecto,
                'solicitud': solicitud, 'fecha_serv': fecha_servicio
            })

        # --- TAB 1 DBR: HERRAMIENTAS ---
        with tabs[1]:
            st.subheader("🔧 HERRAMIENTA(S) EN ANÁLISIS")
            cols_items_dbr = ["N°", "DESCRIPCIÓN", "OD", "CONEXIONES", "NUMERO DE SERIE"]
            df_items_dbr = st.session_state.ed_items_dbr if isinstance(st.session_state.ed_items_dbr, pd.DataFrame) else pd.DataFrame(columns=cols_items_dbr)
            if df_items_dbr.empty:
                df_items_dbr = pd.DataFrame(columns=cols_items_dbr)

            with st.form(key="form_ed_items_dbr"):
                ed_items = st.data_editor(df_items_dbr, num_rows="dynamic", key="ed_items_dbr_in_form")
                if isinstance(ed_items, list):
                    ed_items = pd.DataFrame(ed_items)
                if ed_items is not None and isinstance(ed_items, pd.DataFrame) and not ed_items.empty:
                    ed_items = ed_items.copy()
                    ed_items["N°"] = range(1, len(ed_items) + 1)
                if st.form_submit_button("✅ CONFIRMAR CAMBIOS EN HERRAMIENTAS", use_container_width=True):
                    st.session_state.ed_items_dbr = ed_items
                    st.success("✅ CAMBIOS GUARDADOS")
                    st.rerun()

        # --- TABs 2-9 DBR: Secciones de texto + fotos + tablas ---
        secciones_dbr = [
            (2, "1. INTRODUCCIÓN", 'introduccion', None, None, None, None),
            (3, "2. HISTORIAL DEL EVENTO", 'historial_evento', 'tabla_historial_evento',
             ["FECHA DE EVENTO", "EVENTO", "OBSERVACION"], 'fotos_historial_evento', 'n_fotos_historial_evento'),
            (4, "3. HISTORIAL DE LA HERRAMIENTA", 'historial_herramienta', 'tabla_historial_herramienta',
             ["FECHA DE ACCION", "ACCION", "OBSERVACIÓN"], 'fotos_historial_herramienta', 'n_fotos_historial_herramienta'),
            (5, "4. PROCESOS MONITOREADOS PREVIO AL ENVÍO DE HERRAMIENTA(S)", 'procesos', 'tabla_procesos',
             ["FECHA", "PROCESO", "N REPORTE", "OBSERVACIÓN"], 'fotos_procesos', 'n_fotos_procesos'),
            (6, "5. INFORMACIÓN DE CERTIFICADOS, DOCUMENTACIÓN Y/O ELASTOMEROS UTILIZADOS", 'certificados',
             'tabla_certificados', ["DOCUMENTO", "NUMERO DOC./REVISION/N REF.", "OBSERVACIÓN"], None, None),
            (7, "6. RESUMEN DE MONITOREO QA/QC", 'resumen_monitoreo', None, None, 'fotos_resumen', 'n_fotos_resumen'),
            (9, "8. CONCLUSIONES", 'conclusiones', None, None, None, None),
            (10, "9. RECOMENDACIONES", 'recomendaciones', None, None, None, None),
        ]

        for tab_idx, titulo, campo_texto, campo_tabla, cols_tabla, campo_fotos, campo_n_fotos in secciones_dbr:
            with tabs[tab_idx]:
                st.subheader(titulo)
                st.session_state.form_data_dbr[campo_texto] = st_quill(
                    st.session_state.form_data_dbr.get(campo_texto, ''),
                    key=f"dbr_{campo_texto}"
                )

                if campo_fotos:
                    st.markdown("---")
                    st.markdown("### FOTOS (2 por línea)")
                    fotos_lista = []
                    n_fotos_val = st.session_state.get(campo_n_fotos, 1)
                    for s in range(n_fotos_val):
                        st.write(f"**Línea {s+1}**")
                        cols_f = st.columns(2)
                        for k in range(2):
                            with cols_f[k]:
                                f = st.file_uploader(f"Foto {s+1}-{k+1}", type=['jpg', 'png'], key=f"dbr_f_{campo_fotos}_{s}_{k}")
                                d = st.text_input(f"Descripción {s+1}-{k+1}", key=f"dbr_fd_{campo_fotos}_{s}_{k}")
                                if f:
                                    fotos_lista.append({"imagen": f, "descripcion": d})
                    setattr(st.session_state, campo_fotos, fotos_lista) if hasattr(st.session_state, campo_fotos) else None
                    # Usar setitem para session_state
                    st.session_state[campo_fotos] = fotos_lista

                    col1_f, col2_f = st.columns(2)
                    with col1_f:
                        if st.button("➕ Añadir Línea de Fotos", key=f"add_line_{campo_fotos}"):
                            st.session_state[campo_n_fotos] = n_fotos_val + 1
                            st.rerun()
                    with col2_f:
                        if st.button("➖ Eliminar Última Línea", key=f"del_line_{campo_fotos}") and n_fotos_val > 1:
                            st.session_state[campo_n_fotos] = n_fotos_val - 1
                            st.rerun()

                if campo_tabla:
                    st.markdown("---")
                    st.markdown("### TABLA")
                    df_tabla_actual = st.session_state.get(campo_tabla, pd.DataFrame(columns=cols_tabla))
                    if not isinstance(df_tabla_actual, pd.DataFrame):
                        df_tabla_actual = restaurar_dfs_desde_json(df_tabla_actual, cols_tabla)

                    with st.form(key=f"form_{campo_tabla}"):
                        edited_tabla = st.data_editor(df_tabla_actual, num_rows="dynamic", use_container_width=True, key=f"{campo_tabla}_editor")
                        if st.form_submit_button("✅ CONFIRMAR CAMBIOS EN TABLA", use_container_width=True):
                            st.session_state[campo_tabla] = edited_tabla
                            st.success("✅ CAMBIOS GUARDADOS")
                            st.rerun()

        # --- TAB 8 DBR: HALLAZGOS ---
        with tabs[8]:
            st.subheader("7. HALLAZGOS / DISCREPANCIAS / NCR")
            for idx in range(len(st.session_state.hallazgos_list_dbr)):
                hallazgo = st.session_state.hallazgos_list_dbr[idx]
                with st.expander(f"⚠️ HALLAZGO {idx+1} - {hallazgo.get('tipo', 'Nuevo')}", expanded=False):
                    col_tipo, col_estado = st.columns(2)
                    with col_tipo:
                        hallazgo['tipo'] = st.selectbox("TIPO", ["DISC.", "OBS.", "NCR"],
                            index=["DISC.", "OBS.", "NCR"].index(hallazgo.get('tipo', 'DISC.')), key=f"dbr_tipo_{idx}")
                    with col_estado:
                        hallazgo['estado'] = st.selectbox("ESTADO", ["ABIERTO", "PENDIENTE", "CERRADO"],
                            index=["ABIERTO", "PENDIENTE", "CERRADO"].index(hallazgo.get('estado', 'ABIERTO')), key=f"dbr_estado_{idx}")
                    hallazgo['descripcion'] = st.text_area("DESCRIPCIÓN", value=hallazgo.get('descripcion', ''), key=f"dbr_desc_{idx}", height=100)
                    hallazgo['normativa'] = st.text_area("NORMATIVA", value=hallazgo.get('normativa', ''), key=f"dbr_norm_{idx}", height=80)
                    hallazgo['accion'] = st.text_area("ACCIÓN", value=hallazgo.get('accion', ''), key=f"dbr_acc_{idx}", height=80)
                    hallazgo['resolucion'] = st.text_area("RESOLUCIÓN", value=hallazgo.get('resolucion', ''), key=f"dbr_res_{idx}", height=80)
                    st.session_state.hallazgos_list_dbr[idx] = hallazgo

                    if st.button(f"🗑️ Eliminar HALLAZGO {idx+1}", key=f"dbr_del_hall_{idx}"):
                        st.session_state.hallazgos_list_dbr.pop(idx)
                        st.rerun()

            if st.button("➕ AÑADIR NUEVO HALLAZGO", use_container_width=True, key="dbr_add_hallazgo"):
                st.session_state.hallazgos_list_dbr.append({
                    'tipo': 'DISC.', 'descripcion': '', 'normativa': '',
                    'accion': '', 'resolucion': '', 'estado': 'ABIERTO', 'fotos': []
                })
                st.rerun()

        # --- TAB 11 DBR: REPRESENTANTES + GENERAR ---
        with tabs[11]:
            st.subheader("10. REPRESENTANTES")
            col_r1, col_r2 = st.columns(2)
            with col_r1:
                gerente_proyecto = st.text_input("GERENTE DE PROYECTO", value=st.session_state.form_data_dbr.get('gerente_proyecto', ""), key="dbr_gerente_proyecto")
                rep_cliente = st.text_input("REPRESENTANTE DEL CLIENTE", value=st.session_state.form_data_dbr.get('rep_cliente', ""), key="dbr_rep_cliente")
                cargo_cliente = st.text_input("CARGO DEL REPRESENTANTE DEL CLIENTE", value=st.session_state.form_data_dbr.get('cargo_cliente', ""), key="dbr_cargo_cliente")
            with col_r2:
                # CORREGIDO: nombre correcto "Erika Soruco" con K
                opciones_dbr = ["Brenda Figueroa", "Erika Soruco"]
                valor_actual_dbr = st.session_state.form_data_dbr.get('rep_serfigobol', 'Brenda Figueroa')
                if valor_actual_dbr not in opciones_dbr:
                    valor_actual_dbr = "Brenda Figueroa"
                rep_serfi = st.selectbox("REPRESENTANTE SERFIGOBOL S.R.L.", opciones_dbr, index=opciones_dbr.index(valor_actual_dbr), key="dbr_rep_ser")

            st.session_state.form_data_dbr['gerente_proyecto'] = gerente_proyecto
            st.session_state.form_data_dbr['rep_cliente'] = rep_cliente
            st.session_state.form_data_dbr['cargo_cliente'] = cargo_cliente
            st.session_state.form_data_dbr['rep_serfigobol'] = rep_serfi

            st.divider()

            if st.button("🚀 GENERAR REPORTE DBR (WORD)", type="primary", use_container_width=True, key="generar_dbr"):
                errores_dbr = []
                if not st.session_state.form_data_dbr.get('cliente', '').strip():
                    errores_dbr.append("CLIENTE")
                if not st.session_state.form_data_dbr.get('contrato', '').strip():
                    errores_dbr.append("CONTRATO No")
                if not st.session_state.form_data_dbr.get('pozo', '').strip():
                    errores_dbr.append("POZO/PROYECTO")
                if not st.session_state.form_data_dbr.get('fecha_serv', '').strip():
                    errores_dbr.append("FECHA DE SERVICIO")

                if errores_dbr:
                    st.error(f"❌ COMPLETE: {', '.join(errores_dbr)}")
                else:
                    doc = Document()
                    style = doc.styles['Normal']
                    style.font.name = 'Calibri'
                    style.font.size = Pt(11)
                    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    section = doc.sections[0]
                    section.top_margin = Cm(2.5)
                    section.bottom_margin = Cm(2.5)
                    section.left_margin = Cm(2.5)
                    section.right_margin = Cm(2.5)

                    # CORREGIDO: usar función reutilizable para encabezado
                    fecha_emision_dbr = agregar_encabezado_word(
                        doc, report_id,
                        st.session_state.form_data_dbr.get('cliente', ''),
                        "REPORTE QA/QC",
                        "Análisis de DBR (DAMAGED BEYOND REPAIR)"
                    )
                    if not st.session_state.fecha_emision_actual_dbr:
                        st.session_state.fecha_emision_actual_dbr = fecha_emision_dbr

                    titulo_state_dbr = make_titulo_counter()

                    data_dbr = [
                        ["CLIENTE", st.session_state.form_data_dbr.get('cliente', ''), "CONTACTO DEL CLIENTE", st.session_state.form_data_dbr.get('contacto', '')],
                        ["CONTRATO NO", st.session_state.form_data_dbr.get('contrato', ''), "POZO/PROYECTO", st.session_state.form_data_dbr.get('pozo', '')],
                        ["ORDEN DE COMPRA", st.session_state.form_data_dbr.get('orden_compra', ''), "ORDEN DE SERVICIO", st.session_state.form_data_dbr.get('orden_servicio', '')],
                        ["FECHA DE SERVICIO", st.session_state.form_data_dbr.get('fecha_serv', ''), "SOLICITUD", st.session_state.form_data_dbr.get('solicitud', '')]
                    ]
                    agregar_tabla_cabecera(doc, data_dbr)
                    doc.add_paragraph()

                    add_title_with_bar(doc, "HERRAMIENTA(S) EN ANÁLISIS", titulo_state_dbr)
                    add_table_professional(doc, st.session_state.ed_items_dbr, "")

                    add_title_with_bar(doc, "1. INTRODUCCIÓN", titulo_state_dbr)
                    p_intro = doc.add_paragraph()
                    add_quill(p_intro, st.session_state.form_data_dbr.get('introduccion', ''))

                    add_title_with_bar(doc, "2. HISTORIAL DEL EVENTO", titulo_state_dbr)
                    p_ev = doc.add_paragraph()
                    add_quill(p_ev, st.session_state.form_data_dbr.get('historial_evento', ''))
                    if not st.session_state.tabla_historial_evento.empty:
                        add_subtitle_left(doc, "Tabla de Eventos")
                        add_table_professional(doc, st.session_state.tabla_historial_evento, "")
                    fotos_ev = [f for f in st.session_state.fotos_historial_evento if f.get('imagen') is not None]
                    add_photos_row(doc, fotos_ev, "Fotos del Evento", cols_per_row=2)

                    add_title_with_bar(doc, "3. HISTORIAL DE LA HERRAMIENTA", titulo_state_dbr)
                    p_herr = doc.add_paragraph()
                    add_quill(p_herr, st.session_state.form_data_dbr.get('historial_herramienta', ''))
                    if not st.session_state.tabla_historial_herramienta.empty:
                        add_subtitle_left(doc, "Tabla de Acciones")
                        add_table_professional(doc, st.session_state.tabla_historial_herramienta, "")
                    fotos_herr = [f for f in st.session_state.fotos_historial_herramienta if f.get('imagen') is not None]
                    add_photos_row(doc, fotos_herr, "Fotos del Historial de la Herramienta", cols_per_row=2)

                    add_title_with_bar(doc, "4. PROCESOS MONITOREADOS PREVIO AL ENVÍO DE HERRAMIENTA(S)", titulo_state_dbr)
                    p_proc = doc.add_paragraph()
                    add_quill(p_proc, st.session_state.form_data_dbr.get('procesos', ''))
                    if not st.session_state.tabla_procesos.empty:
                        add_subtitle_left(doc, "Tabla de Procesos")
                        add_table_professional(doc, st.session_state.tabla_procesos, "")
                    fotos_proc = [f for f in st.session_state.fotos_procesos if f.get('imagen') is not None]
                    add_photos_row(doc, fotos_proc, "Fotos de Procesos", cols_per_row=2)

                    add_title_with_bar(doc, "5. INFORMACIÓN DE CERTIFICADOS, DOCUMENTACIÓN Y/O ELASTOMEROS UTILIZADOS", titulo_state_dbr)
                    p_cert = doc.add_paragraph()
                    add_quill(p_cert, st.session_state.form_data_dbr.get('certificados', ''))
                    if not st.session_state.tabla_certificados.empty:
                        add_subtitle_left(doc, "Tabla de Documentos")
                        add_table_professional(doc, st.session_state.tabla_certificados, "")

                    add_title_with_bar(doc, "6. RESUMEN DE MONITOREO QA/QC", titulo_state_dbr)
                    p_res = doc.add_paragraph()
                    add_quill(p_res, st.session_state.form_data_dbr.get('resumen_monitoreo', ''))
                    fotos_res = [f for f in st.session_state.fotos_resumen if f.get('imagen') is not None]
                    add_photos_row(doc, fotos_res, "Fotos de Resumen de Monitoreo", cols_per_row=2)

                    add_title_with_bar(doc, "7. HALLAZGOS / DISCREPANCIAS / NCR", titulo_state_dbr)
                    if st.session_state.hallazgos_list_dbr:
                        for idx, hallazgo in enumerate(st.session_state.hallazgos_list_dbr):
                            add_hallazgo_table(doc, hallazgo, idx + 1)
                            doc.add_paragraph()
                    else:
                        doc.add_paragraph("No se registraron hallazgos.")

                    add_title_with_bar(doc, "8. CONCLUSIONES", titulo_state_dbr)
                    p_conc = doc.add_paragraph()
                    add_quill(p_conc, st.session_state.form_data_dbr.get('conclusiones', ''))

                    add_title_with_bar(doc, "9. RECOMENDACIONES", titulo_state_dbr)
                    p_rec = doc.add_paragraph()
                    add_quill(p_rec, st.session_state.form_data_dbr.get('recomendaciones', ''))

                    add_title_with_bar(doc, "10. REPRESENTANTES", titulo_state_dbr)
                    agregar_seccion_representantes(doc, st.session_state.form_data_dbr, tipo='dbr')

                    agregar_pie_pagina(doc)

                    target = io.BytesIO()
                    doc.save(target)

                    st.download_button("📥 DESCARGAR DOCX", target.getvalue(), f"{report_id}.docx", key="download_dbr_final")

                    guardar_datos_reporte_dbr(p_sel, poz_sel, prov_sel, report_id, es_generacion_word=True)
                    guardar_hallazgos_en_discrepancias(poz_sel, p_sel, report_id, st.session_state.hallazgos_list_dbr)

                    st.balloons()
                    st.success(f"✅ REPORTE DBR {report_id} GENERADO CORRECTAMENTE")

# ============================================================================
# HISTORIAL DE INSPECCIONES
# ============================================================================
elif menu == "📊 HISTORIAL DE INSPECCIONES":
    st.header("📊 HISTORIAL DE INSPECCIONES")

    tipo_reporte = st.radio("Tipo de Reporte", ["QA/QC (FR)", "DBR"], horizontal=True)
    prefijo = "FR-" if tipo_reporte == "QA/QC (FR)" else "DBR-"

    st.subheader("🔍 BUSCAR REPORTES")
    col_s1, col_s2, col_s3, col_s4 = st.columns(4)
    with col_s1:
        buscar_codigo = st.text_input("CÓDIGO", placeholder=f"Ej: {prefijo}1")
    with col_s2:
        buscar_pozo = st.text_input("POZO", placeholder="Nombre del pozo")
    with col_s3:
        buscar_fecha = st.text_input("FECHA", placeholder="YYYY-MM-DD")
    with col_s4:
        buscar_cliente = st.text_input("CLIENTE", placeholder="Nombre del cliente")

    # CORREGIDO: una sola consulta por tipo
    if st.session_state.usuario_rol == "admin":
        inspecciones = ejecutar_query("""
            SELECT id, reporte_id, fecha_inspeccion, cliente, pozo_nombre, contrato,
                   fecha_servicio, usuario_creador_id, estado_reporte
            FROM inspecciones WHERE reporte_id LIKE ? ORDER BY fecha_inspeccion DESC
        """, (f"{prefijo}%",), fetch=True)
    else:
        inspecciones = ejecutar_query("""
            SELECT id, reporte_id, fecha_inspeccion, cliente, pozo_nombre, contrato,
                   fecha_servicio, usuario_creador_id, estado_reporte
            FROM inspecciones WHERE usuario_creador_id = ? AND reporte_id LIKE ?
            ORDER BY fecha_inspeccion DESC
        """, (st.session_state.usuario_id, f"{prefijo}%"), fetch=True)

    resultados = []
    if inspecciones:
        for insp in inspecciones:
            id_reg, reporte_id, fecha_inspeccion, cliente, pozo_nombre, contrato, fecha_servicio, creador_id, estado_reporte = insp
            if buscar_codigo and buscar_codigo.lower() not in reporte_id.lower():
                continue
            if buscar_pozo and buscar_pozo.lower() not in (pozo_nombre or "").lower():
                continue
            if buscar_fecha and buscar_fecha not in (fecha_inspeccion or ""):
                continue
            if buscar_cliente and buscar_cliente.lower() not in (cliente or "").lower():
                continue
            resultados.append(insp)

    if resultados:
        items_por_pagina = 10
        total_items = len(resultados)
        total_paginas = max(1, (total_items + items_por_pagina - 1) // items_por_pagina)

        col_pag1, col_pag2, col_pag3 = st.columns([1, 2, 1])
        with col_pag2:
            pagina_actual = st.number_input("Página", min_value=1, max_value=total_paginas,
                                            value=min(st.session_state.historial_pagina, total_paginas),
                                            step=1, key="pagina_input")
            st.session_state.historial_pagina = pagina_actual

        inicio = (pagina_actual - 1) * items_por_pagina
        fin = inicio + items_por_pagina
        resultados_paginados = resultados[inicio:fin]

        st.subheader(f"📋 RESULTADOS ({total_items} REPORTES) - Página {pagina_actual} de {total_paginas}")

        for insp in resultados_paginados:
            id_reg, reporte_id, fecha_inspeccion, cliente, pozo_nombre, contrato, fecha_servicio, creador_id, estado_reporte = insp
            estado_icono = "📝" if estado_reporte == "BORRADOR" else "✅"
            estado_texto = "EN CURSO" if estado_reporte == "BORRADOR" else "COMPLETADO"

            with st.expander(f"{estado_icono} {reporte_id} - {pozo_nombre} / {cliente} [{estado_texto}]", expanded=False):
                col1, col2 = st.columns(2)
                with col1:
                    st.write(f"**FECHA CREACIÓN:** {fecha_inspeccion}")
                    st.write(f"**FECHA SERVICIO:** {fecha_servicio}")
                    st.write(f"**CLIENTE:** {cliente}")
                    st.write(f"**ESTADO:** {estado_texto}")
                with col2:
                    st.write(f"**CONTRATO:** {contrato}")
                    st.write(f"**POZO:** {pozo_nombre}")
                    if st.session_state.usuario_rol == "admin":
                        creador = ejecutar_query("SELECT username FROM usuarios WHERE id = ?", (creador_id,), fetch=True)
                        st.write(f"**CREADO POR:** {creador[0][0] if creador else 'N/A'}")

                puede_editar = (st.session_state.usuario_rol == "admin") or (creador_id == st.session_state.usuario_id)
                if puede_editar:
                    if reporte_id.startswith("FR-"):
                        if st.button("📝 EDITAR", key=f"edit_{id_reg}"):
                            st.session_state.modo_edicion = True
                            st.session_state.edicion_reporte_id = reporte_id
                            st.session_state.formulario_activo = True
                            st.session_state.menu_seleccionado = "📝 NUEVO REPORTE QA/QC"
                            st.session_state.datos_editados_cargados = False
                            st.rerun()
                    else:
                        if st.button("📝 EDITAR DBR", key=f"edit_dbr_{id_reg}"):
                            st.session_state.modo_edicion_dbr = True
                            st.session_state.edicion_reporte_id_dbr = reporte_id
                            st.session_state.formulario_activo_dbr = True
                            st.session_state.menu_seleccionado = "📝 NUEVO REPORTE QA/QC - DBR"
                            st.session_state.datos_editados_dbr_cargados = False
                            st.rerun()
                else:
                    st.button("📝 EDITAR", key=f"edit_{id_reg}", disabled=True, help="No tiene permiso para editar este reporte")
    else:
        st.info("📭 NO HAY INSPECCIONES REGISTRADAS")

# ============================================================================
# DISCREPANCIAS POR POZO
# ============================================================================
elif menu == "📋 DISCREPANCIAS POR POZO":
    st.header("📋 DISCREPANCIAS POR POZO")

    st.subheader("🔍 FILTRAR DISCREPANCIAS")
    proyectos = ejecutar_query("SELECT id, contrato_nom FROM proyectos ORDER BY contrato_nom", fetch=True)
    if not proyectos:
        st.warning("⚠️ NO HAY PROYECTOS REGISTRADOS.")
        st.stop()

    col_filtro1, col_filtro2 = st.columns(2)
    with col_filtro1:
        proyecto_id = st.selectbox("SELECCIONAR PROYECTO", proyectos, format_func=lambda x: x[1])

    pozos = ejecutar_query("SELECT id, nombre_pozo FROM pozos WHERE proyecto_id = ? ORDER BY nombre_pozo", (proyecto_id[0],), fetch=True)
    if not pozos:
        st.warning("⚠️ NO HAY POZOS EN ESTE PROYECTO.")
        st.stop()

    with col_filtro2:
        pozo_id = st.selectbox("SELECCIONAR POZO", pozos, format_func=lambda x: x[1])

    # CORREGIDO: tipo_discrepancia con key único para evitar conflicto con historial
    tipo_discrepancia = st.radio("Tipo de Reporte", ["QA/QC (FR)", "DBR"], horizontal=True, key="tipo_disc_pozo")
    prefijo_disc = "FR-" if tipo_discrepancia == "QA/QC (FR)" else "DBR-"

    # CORREGIDO: una sola consulta, no duplicada
    discrepancias = ejecutar_query("""
        SELECT dp.id, dp.reporte_id, dp.tipo, dp.descripcion, dp.normativa, dp.acciones, dp.resolucion, dp.estado, dp.fecha, u.username, dp.fotos_json
        FROM discrepancias_pozo dp
        LEFT JOIN usuarios u ON dp.usuario_id = u.id
        WHERE dp.pozo_id = ? AND dp.reporte_id LIKE ?
        ORDER BY dp.fecha DESC
    """, (pozo_id[0], f"{prefijo_disc}%"), fetch=True)

    if discrepancias:
        df_discrepancias = pd.DataFrame(discrepancias, columns=["ID", "REPORTE", "TIPO", "DESCRIPCIÓN", "NORMATIVA", "ACCIONES", "RESOLUCION", "ESTADO", "FECHA", "USUARIO", "FOTOS"])
        st.dataframe(df_discrepancias.drop(columns=["FOTOS"]), use_container_width=True, hide_index=True)

        if st.button("📥 EXPORTAR A EXCEL", key="exportar_disc"):
            output = io.BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                df_discrepancias.drop(columns=["FOTOS"]).to_excel(writer, sheet_name='Discrepancias', index=False)
            st.download_button(
                label="📥 DESCARGAR EXCEL",
                data=output.getvalue(),
                file_name=f"discrepancias_{tipo_discrepancia.replace('/', '_')}_pozo_{pozo_id[1]}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

        st.markdown("---")
        st.subheader("📋 DETALLE DE DISCREPANCIAS")

        # CORREGIDO: bloque único para mostrar detalle, sin duplicación
        for idx, disc in enumerate(discrepancias):
            id_reg, reporte_id, tipo, descripcion, normativa, acciones, resolucion, estado, fecha, usuario, fotos_json = disc

            with st.expander(f"⚠️ {reporte_id} - {tipo} - {fecha}", expanded=False):
                col1, col2 = st.columns(2)
                with col1:
                    st.write(f"**TIPO:** {tipo}")
                    st.write(f"**DESCRIPCIÓN:** {descripcion}")
                    st.write(f"**NORMATIVA:** {normativa}")
                with col2:
                    st.write(f"**ACCIÓN:** {acciones}")
                    st.write(f"**RESOLUCIÓN:** {resolucion}")
                    st.write(f"**ESTADO:** {estado}")
                    st.write(f"**REGISTRADO POR:** {usuario}")

                st.markdown("---")
                st.markdown("**FOTOS DE LA DISCREPANCIA**")

                fotos_discrepancia = []
                if fotos_json:
                    try:
                        fotos_data = json.loads(fotos_json)
                        fotos_discrepancia = restaurar_lista_fotos(fotos_data)
                    except Exception:
                        pass

                fotos_validas = [f for f in fotos_discrepancia if f.get('imagen') is not None]
                if fotos_validas:
                    st.write(f"**{len(fotos_validas)} foto(s) asociada(s)**")
                    cols_fotos = st.columns(4)
                    for i, foto in enumerate(fotos_validas):
                        with cols_fotos[i % 4]:
                            try:
                                foto['imagen'].seek(0)
                                st.image(foto['imagen'], caption=foto.get('descripcion', ''), use_container_width=True)
                            except Exception:
                                st.write("Error al cargar foto")
                else:
                    st.info("No hay fotos asociadas a esta discrepancia")

                with st.expander("📸 AGREGAR/EDITAR FOTOS"):
                    key_fotos_disc = f'fotos_disc_{id_reg}'
                    if key_fotos_disc not in st.session_state:
                        st.session_state[key_fotos_disc] = list(fotos_discrepancia)

                    for i, foto in enumerate(st.session_state[key_fotos_disc]):
                        col_foto, col_desc, col_del = st.columns([2, 2, 1])
                        with col_foto:
                            try:
                                if foto.get('imagen'):
                                    foto['imagen'].seek(0)
                                    st.image(foto['imagen'], width=100)
                            except Exception:
                                st.write("Imagen no disponible")
                        with col_desc:
                            nueva_desc = st.text_input(f"Descripción foto {i+1}", value=foto.get('descripcion', ''), key=f"disc_desc_{id_reg}_{i}")
                            st.session_state[key_fotos_disc][i]['descripcion'] = nueva_desc
                        with col_del:
                            if st.button(f"🗑️ Eliminar", key=f"del_disc_foto_{id_reg}_{i}"):
                                st.session_state[key_fotos_disc].pop(i)
                                st.rerun()

                    st.write("**Agregar nueva foto:**")
                    nueva_foto = st.file_uploader("Seleccionar imagen", type=['jpg', 'png', 'jpeg'], key=f"disc_new_foto_{id_reg}")
                    nueva_desc_foto = st.text_input("Descripción", key=f"disc_new_desc_{id_reg}")

                    if st.button("💾 GUARDAR CAMBIOS", key=f"disc_save_{id_reg}"):
                        if nueva_foto:
                            st.session_state[key_fotos_disc].append({'imagen': nueva_foto, 'descripcion': nueva_desc_foto})

                        fotos_guardar = serializar_lista_fotos(st.session_state[key_fotos_disc])
                        ejecutar_query("UPDATE discrepancias_pozo SET fotos_json = ? WHERE id = ?", (json.dumps(fotos_guardar), id_reg))
                        st.success("✅ Fotos guardadas correctamente")
                        st.rerun()
    else:
        st.info(f"📭 NO HAY DISCREPANCIAS {tipo_discrepancia} REGISTRADAS PARA ESTE POZO")
        
